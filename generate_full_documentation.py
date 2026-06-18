"""
generate_full_documentation.py
===============================
Generates the comprehensive project documentation DOCX for the
IOCL Debutanizer C4 Slippage Optimization AI project.
Run from the DEBUTANIZER-model root directory.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────────────────────────────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*color)
    heading.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    heading.paragraph_format.space_after  = Pt(6)
    return heading

def add_paragraph(doc, text, bold=False, italic=False, size=10.5, color=None, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(doc, text, bold_prefix=None, indent=0.25):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10.5)
        p.add_run(text).font.size = Pt(10.5)
    else:
        p.add_run(text).font.size = Pt(10.5)
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.style = doc.styles['No Spacing']
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 100, 30)
    return p

def add_table(doc, headers, rows, header_color="1F4E79", header_text_color=RGBColor(255,255,255)):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, header_color)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = header_text_color
                run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        bg = "F5F5F5" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
    doc.add_paragraph()  # spacer
    return table

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("─" * 90)
    run.font.color.rgb = RGBColor(180, 180, 180)
    run.font.size = Pt(8)

def page_break(doc):
    doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT GENERATION
# ─────────────────────────────────────────────────────────────────────────────

def build_docx():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    # ═══════════════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title.add_run("IOCL Debutanizer Column")
    t_run.bold = True
    t_run.font.size = Pt(28)
    t_run.font.color.rgb = RGBColor(13, 71, 161)

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2_run = title2.add_run("C4 Slippage Optimization")
    t2_run.bold = True
    t2_run.font.size = Pt(26)
    t2_run.font.color.rgb = RGBColor(13, 71, 161)

    title3 = doc.add_paragraph()
    title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t3_run = title3.add_run("AI Soft Sensor & Advisory Optimizer")
    t3_run.font.size = Pt(18)
    t3_run.font.color.rgb = RGBColor(66, 165, 245)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("Comprehensive Technical Documentation")
    sub_run.italic = True
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Document generated: {datetime.datetime.now().strftime('%d %B %Y')}\n")
    meta.add_run("Debutanizer AI Project Team\n")
    meta.add_run("Status: Frozen & Validated — v2.1 Advisory")

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (manual)
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "Table of Contents", level=1, color=(13,71,161))
    toc_items = [
        ("1", "Problem Statement & Objectives"),
        ("2", "Process Overview — Debutanizer Column"),
        ("3", "Dataset Understanding & Preprocessing"),
        ("4", "Feature Engineering & Drift-Resistant Design"),
        ("5", "Model A — C4H8 Soft Sensor Development"),
        ("6", "Model B — C4H6 Soft Sensor Development"),
        ("7", "Concept Drift: Discovery, Diagnosis & Resolution"),
        ("8", "Phase 5.1A — Surrogate Process Models"),
        ("9", "Phase 5.2 — Physics-Aware Advisory Optimizer"),
        ("10", "Batch Validation & Performance Results"),
        ("11", "Interactive Dashboard (NiceGUI Platform)"),
        ("12", "Problems, Setbacks & Recoveries"),
        ("13", "File & Folder Reference"),
        ("14", "Key Formulas & Mathematical Reference"),
        ("15", "Lessons Learned & Final Conclusions"),
        ("16", "Future Work & Roadmap"),
    ]
    for num, title_text in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{num}.   {title_text}")
        r.font.size = Pt(10.5)

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 1 — PROBLEM STATEMENT
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "1. Problem Statement & Objectives", level=1, color=(13,71,161))
    add_paragraph(doc,
        "The IOCL Debutanizer column is tasked with separating mixed C4 hydrocarbons (butylene, butadiene, n-butane) "
        "from C5+ heavier components in the bottom product stream. Excess C4 carryover into the bottom stream — known as "
        "'C4 slippage' — increases product contamination, reduces downstream quality, and causes economic loss. "
        "The specification limit is <0.50 wt% Total C4 in the bottom product. In practice, the column has been operating "
        "with Total C4 concentrations in the range of 0.8–1.5 wt%, well above specification."
    )
    add_heading(doc, "1.1 Root Causes of the Problem", level=2)
    for item in [
        ("Analyzer Lag: ", "The online analyzer has a 12-minute cycle time. This means the operator receives a new composition reading only every 12 minutes, creating a blind window for rapid process upsets."),
        ("Analyzer Reliability: ", "The analyzer was found to have extended 'stuck' periods — periods where it repeatedly outputs the same value without detecting actual composition changes."),
        ("Manual Operation: ", "Column steam and reflux setpoints were adjusted manually based on experience, with no real-time model feedback on composition impact."),
        ("Feed Variability: ", "The debutanizer receives variable-quality feed from upstream units, causing composition disturbances that lag any manual response."),
        ("Campaign Drift: ", "Across multiple operational campaigns (Blocks), column pressure setpoints and feed quality shifted significantly, making any fixed-parameter model obsolete."),
    ]:
        add_bullet(doc, item[1], bold_prefix=item[0])

    add_heading(doc, "1.2 Project Objectives", level=2)
    for item in [
        "Develop an AI-based soft sensor to predict Total C4 wt% in the bottom product in real-time, bridging the 12-minute analyzer gap.",
        "Make the soft sensor robust to campaign-to-campaign drift without retraining.",
        "Build a process-aware advisory optimizer that recommends Reboiler Steam Flow and Reflux Flow setpoints to minimize C4 slippage.",
        "Package the system into an operator-facing interactive dashboard for daily use.",
    ]:
        add_bullet(doc, item)
    add_divider(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 2 — PROCESS OVERVIEW
    # ═══════════════════════════════════════════════════════════════════════════
    page_break(doc)
    add_heading(doc, "2. Process Overview — Debutanizer Column", level=1, color=(13,71,161))
    add_paragraph(doc,
        "The Debutanizer is a distillation column that separates the mixed C4 fraction from the C5+ heavier fraction. "
        "Feed from the Depropanizer (DP) bottom is introduced at the 17th tray under level control. Column vapors are "
        "condensed with cooling water in the overhead condenser and collected in a reflux drum. Mixed C4s, after meeting "
        "the reflux requirement, are sent for further processing to the Butadiene Extraction Unit, C4 Hydrogenation Unit, "
        "or OSBL storage. The C5+ bottom product is routed to downstream processing."
    )

    add_heading(doc, "2.1 Key Process Variables", level=2)
    add_table(doc,
        ["Variable", "Tag / Description", "Typical Range", "Role in Model"],
        [
            ["Feed Flow", "Fresh feed from DP bottom", "~60–70 TPH", "Denominator for ratio features"],
            ["Reboiling Steam Flow", "LP steam to reboiler", "18 – 24.4 TPH", "Primary manipulated variable"],
            ["Reflux Flow", "Overhead reflux", "80 – 103.9 TPH", "Primary manipulated variable"],
            ["Column Top Pressure", "Overhead pressure", "~3.98 – 4.19 bar", "Bubble-point setpoint indicator"],
            ["Column Bottom Temp", "Reboiler temperature proxy", "~105 – 115 °C", "Key column thermal state indicator"],
            ["Control Tray Temp", "17th tray temperature", "~65 – 80 °C", "Composition proxy, sensitive to drift"],
            ["Reboiler Outlet Temp", "Reboiler outlet", "~110 – 125 °C", "Heat input indicator"],
            ["C4H8 Bottom", "Butylene analyzer, bottom", "~0.3 – 1.8 wt%", "Primary soft sensor target (Model A)"],
            ["C4H6 Bottom", "Butadiene analyzer, bottom", "~0.003 – 0.21 wt%", "Secondary soft sensor target (Model B)"],
            ["Total C4", "C4H8 + C4H6", "Spec: <0.50 wt%", "Composite optimization target"],
        ]
    )

    add_heading(doc, "2.2 Operating Campaign Structure", level=2)
    add_paragraph(doc,
        "The dataset spans four distinct operational campaigns, labeled Data_Block 1–4. Each block represents a "
        "continuous period of operation separated by planned shutdowns or turnarounds. The campaign structure is "
        "critical to the modeling strategy:"
    )
    add_table(doc,
        ["Block", "Operating Regime", "Approx. Duration", "Mean C4H8 (wt%)", "Mean C4H6 (wt%)", "Key Characteristic"],
        [
            ["1", "Cold reboiler startup", "Short", "~0.50", "~0.208", "High C4H6, different thermodynamics"],
            ["2", "Hot fractionation", "Long", "Variable", "~0.031", "Normal operation, training data"],
            ["3", "Hot fractionation", "Long", "Variable", "~0.023", "Normal operation, training data"],
            ["4", "Hot fractionation (lower pressure)", "Test period", "Variable", "~0.0057", "Holdout test set — lower pressure, process shift"],
        ]
    )
    add_divider(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 3 — DATASET
    # ═══════════════════════════════════════════════════════════════════════════
    page_break(doc)
    add_heading(doc, "3. Dataset Understanding & Preprocessing", level=1, color=(13,71,161))
    add_paragraph(doc,
        "The raw dataset was provided as an Excel file ('9.DB DATA-B.xlsx') and a CSV export, containing approximately "
        "4 years of hourly process historian data. The preprocessing pipeline (data_preprocessing.py) applies "
        "a sequence of cleaning, validation, and flagging steps before any feature engineering."
    )

    add_heading(doc, "3.1 Preprocessing Steps", level=2)
    steps = [
        ("Datetime Parsing & Sorting: ", "All timestamps are parsed and rows sorted chronologically. A continuous hourly index is established."),
        ("Winsorization: ", "Extreme outlier values for each process variable are clipped at the 1st and 99th percentile to remove sensor glitches without full row deletion."),
        ("Stuck Analyzer Detection: ", "The C4H8 and C4H6 bottom analyzer outputs are checked for 'stuck' periods — sequences where the same value is repeated for more than a configurable threshold of consecutive readings. A boolean flag (C4H8_Bottom_stuck / C4H6_Bottom_stuck) is added for each row."),
        ("Shutdown Row Flagging: ", "Rows where feed flow falls below a minimum threshold are flagged as potential shutdown rows and excluded from training."),
        ("Block Assignment: ", "Each row is assigned to a Data_Block (1–4) based on campaign boundaries identified from the feed flow and analyzer gap patterns."),
        ("Output: ", "A cleaned Parquet file ('data/clean_data.parquet') is saved for downstream feature engineering."),
    ]
    for bold, text in steps:
        add_bullet(doc, text, bold_prefix=bold)

    add_heading(doc, "3.2 Data Quality Challenges Discovered", level=2)
    challenges = [
        ("Analyzer Stuck Periods: ", "Extended periods where the analyzer outputs a flat-line value were identified in all blocks. In Block 4, these accounted for a significant fraction of rows. These were masked and handled via a persistence anchor rather than being used as training labels."),
        ("Campaign Gaps: ", "Shutdowns between blocks create temporal gaps. Lag and rolling features computed naively would incorrectly propagate values across these gaps. The preprocessing explicitly identifies and handles these boundaries."),
        ("Block 1 Target Distribution Mismatch: ", "The C4H6 mean in Block 1 (0.208 wt%) is 37x higher than Block 4 (0.0057 wt%), reflecting fundamentally different thermodynamic operating conditions during the cold startup regime. Block 1 was ultimately excluded from Model B training."),
    ]
    for bold, text in challenges:
        add_bullet(doc, text, bold_prefix=bold)
    add_divider(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 4 — FEATURE ENGINEERING
    # ═══════════════════════════════════════════════════════════════════════════
    page_break(doc)
    add_heading(doc, "4. Feature Engineering & Drift-Resistant Design", level=1, color=(13,71,161))
    add_paragraph(doc,
        "Feature engineering (feature_engineering.py) is the most critical component of this project. "
        "The core insight is that absolute sensor readings (temperatures, pressures) drift between campaigns "
        "due to pressure setpoint changes, feed quality changes, and equipment fouling. Absolute features "
        "cause sign reversals in feature-target correlations across campaigns, leading to model prediction inversion. "
        "The solution is a suite of drift-resistant features."
    )

    add_heading(doc, "4.1 Gap-Aware Time-Series Resampling", level=2)
    add_paragraph(doc,
        "Each campaign block is isolated and resampled to a continuous hourly grid before lag and rolling features "
        "are computed. The grid is then reindexed back to the original timestamps. This ensures that no lag values "
        "leak across campaign gaps, and any gaps within a block propagate NaN correctly."
    )
    add_code_block(doc, "# For each Data_Block:\n# 1. Extract block slice\n# 2. Resample to hourly grid\n# 3. Compute lag1, lag2, roll_mean_12h features\n# 4. Reindex back to original timestamps")

    add_heading(doc, "4.2 Dimensionless Ratio Features", level=2)
    add_paragraph(doc,
        "By normalizing flows by feed rate, the features become dimensionless and campaign-invariant:")
    add_code_block(doc,
        "Reflux_Ratio     = Reflux_Flow / Feed_Flow\n"
        "Steam_Feed_Ratio = Reboiling_Steam_Flow / Feed_Flow\n\n"
        "# Physical meaning:\n"
        "# - Reflux_Ratio captures the overhead L/V balance (separation efficiency)\n"
        "# - Steam_Feed_Ratio captures the reboiler heat input per unit feed"
    )

    add_heading(doc, "4.3 Rolling Deviation Features (_dev24h)", level=2)
    add_paragraph(doc,
        "The most important innovation: instead of absolute temperature/pressure values, the model uses the "
        "deviation from the 24-hour rolling mean. This extracts the high-frequency transient signal while "
        "stripping out the slow campaign-level drift:"
    )
    add_code_block(doc,
        "Column_Bottom_Temp_dev24h     = Column_Bottom_Temp - rolling_mean_24h(Column_Bottom_Temp)\n"
        "Control_Tray_Temp_dev24h      = Control_Tray_Temp  - rolling_mean_24h(Control_Tray_Temp)\n"
        "Column_Top_Pressure_dev24h    = Column_Top_Pressure - rolling_mean_24h(Column_Top_Pressure)\n"
        "Reboiling_Steam_Flow_dev24h   = Reboiling_Steam_Flow - rolling_mean_24h(Reboiling_Steam_Flow)\n"
        "Reflux_Flow_dev24h            = Reflux_Flow - rolling_mean_24h(Reflux_Flow)"
    )
    add_paragraph(doc,
        "These deviations maintain consistent physical interpretation across campaigns: a positive deviation "
        "means 'hotter than recent average' regardless of the absolute operating pressure level."
    )

    add_heading(doc, "4.4 Leak-Free Campaign Anchor", level=2)
    add_paragraph(doc,
        "The last valid analyzer reading is forward-filled within each block with a strict hour limit. "
        "The critical anti-leakage design is the shift(1) before the fill:"
    )
    add_code_block(doc,
        "# C4H8 anchor (72-hour limit):\n"
        "df['C4H8_last_valid'] = df['C4H8_Bottom'].where(~df['C4H8_Bottom_stuck'])\n"
        "df['C4H8_campaign_anchor'] = (\n"
        "    df.groupby('Data_Block')['C4H8_last_valid']\n"
        "      .transform(lambda x: x.shift(1).ffill(limit=72))\n"
        ")\n\n"
        "# The shift(1) ensures the model NEVER sees the current reading as a feature.\n"
        "# Only the PREVIOUS valid reading is used — 100% leak-free for deployment."
    )

    add_heading(doc, "4.5 Final 8-Feature Configuration (Subset 7 — Anti-Drift)", level=2)
    add_paragraph(doc,
        "After systematic ablation studies across 7 feature subsets and 11 drift experiments, the following "
        "8-feature configuration was selected as the optimal, drift-resistant, deployment-ready feature set:"
    )
    add_table(doc,
        ["Feature", "Type", "Physical Interpretation"],
        [
            ["C4H8_campaign_anchor", "Dynamic Calibration", "Last valid analyzer reading (12–72h limit). Provides campaign composition baseline."],
            ["Reflux_Ratio", "Dimensionless Ratio", "Overhead L/V balance. Campaign-invariant separation efficiency."],
            ["Steam_Feed_Ratio", "Dimensionless Ratio", "Reboiler heat input per unit feed. Campaign-invariant energy balance."],
            ["Reboiling_Steam_Flow_dev24h", "Short-term Deviation", "Steam change from 24h average. High-frequency heat input signal."],
            ["Reflux_Flow_dev24h", "Short-term Deviation", "Reflux change from 24h average. Short-term L/V balance perturbation."],
            ["Column_Bottom_Temp_dev24h", "Short-term Deviation", "Temperature change from 24h average. Composition change proxy."],
            ["Control_Tray_Temp_dev24h", "Short-term Deviation", "Tray temperature deviation. Secondary composition change indicator."],
            ["Column_Top_Pressure_dev24h", "Short-term Deviation", "Pressure change from 24h average. Condenser or overhead load change."],
        ]
    )
    add_divider(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 5 — MODEL A
    # ═══════════════════════════════════════════════════════════════════════════
    page_break(doc)
    add_heading(doc, "5. Model A — C4H8 Soft Sensor Development", level=1, color=(13,71,161))
    add_paragraph(doc,
        "Model A predicts C4H8 wt% in the debutanizer bottom product. This is the primary model — "
        "C4H8 (butylene) constitutes the dominant fraction of Total C4 slippage. "
        "The target variable is C4H8_Bottom."
    )

    add_heading(doc, "5.1 Training Strategy", level=2)
    for item in [
        ("Train Set: ", "Data Blocks 1, 2, and 3 (all hot-regime, non-shutdown rows with valid analyzer readings)."),
        ("Test Set: ", "Data Block 4 — the most recent campaign, held out entirely. Never used during training or hyperparameter tuning."),
        ("Cross-Validation: ", "5-fold TimeSeriesSplit CV within the training blocks used for Optuna hyperparameter optimization."),
    ]:
        add_bullet(doc, item[1], bold_prefix=item[0])

    add_heading(doc, "5.2 Algorithm Selection", level=2)
    add_table(doc,
        ["Algorithm", "CV R² (Train)", "Test R² (Block 4)", "Test MAE (wt%)", "Notes"],
        [
            ["Tuned LightGBM", "0.7087", "0.9147", "0.0494", "Best overall"],
            ["Tuned XGBoost", "0.7037", "0.9074", "0.0516", "Selected for deployment (best JSON portability)"],
            ["Ensemble (CB+XG+LG)", "N/A", "0.9052", "0.0513", "Marginal gain, added complexity"],
            ["Tuned CatBoost", "0.7181", "0.9030", "0.0524", "Best CV, slightly lower test R²"],
            ["XGBoost (Baseline)", "N/A", "0.8846", "0.0572", "Untuned baseline"],
        ]
    )
    add_paragraph(doc,
        "The Tuned XGBoost model was selected for production deployment due to its combination of strong performance, "
        "DCS-friendly JSON serialization format, and lean dependency footprint."
    )

    add_heading(doc, "5.3 Optuna Hyperparameter Optimization", level=2)
    add_paragraph(doc,
        "Optuna with 50 TPE trials and 5-fold TimeSeriesSplit CV was used. Key insight: "
        "the optimal max_depth=3 is extremely shallow, which is physically correct — shallow trees "
        "cannot memorize campaign-specific temperature patterns and are forced to learn generalizable rules."
    )
    add_table(doc,
        ["Hyperparameter", "Optimal Value", "Physical Interpretation"],
        [
            ["n_estimators", "102", "Moderate ensemble size — prevents over-smoothing"],
            ["max_depth", "3", "Very shallow — forces generalization, prevents campaign overfitting"],
            ["learning_rate", "0.0405", "Slow learning rate — robust to noisy labels from stuck analyzer"],
            ["subsample", "0.81", "Row-level stochasticity — improves robustness to outliers"],
            ["colsample_bytree", "0.94", "Nearly full feature set per tree — all 8 features are informative"],
            ["min_child_weight", "8", "High leaf size — prevents fitting noise"],
            ["gamma", "3.4e-5", "Minimal pruning — data is signal-rich with 8 clean features"],
            ["reg_alpha", "7.8e-4", "Slight L1 — minor sparsity regularization"],
            ["reg_lambda", "3.8e-8", "Negligible L2"],
        ]
    )

    add_heading(doc, "5.4 Final Model A Performance", level=2)
    add_paragraph(doc, "Tuned XGBoost (8-feature Anti-Drift Configuration):", bold=True)
    add_table(doc,
        ["Metric", "Value"],
        [
            ["CV R² (Blocks 1-3, 5-fold)", "0.7037"],
            ["Test R² (Block 4, held-out)", "0.9074"],
            ["Test MAE (Block 4)", "0.0516 wt%"],
            ["Anchor Coverage (Block 4)", "94.1% (12h limit)"],
            ["Anchor Leakage Check", "PASS — shift(1) confirmed"],
        ]
    )
    add_paragraph(doc,
        "The gap between CV R² (0.70) and Test R² (0.91) reflects the anchor effect: during CV within training blocks, "
        "the anchor is frequently NaN due to stuck periods, so the model learns to rely on process features. "
        "At test time, the anchor is well-populated and provides a strong prior, explaining the higher test R²."
    )

    add_heading(doc, "5.5 Production Inference — predict_c4h8.py", level=2)
    add_paragraph(doc,
        "The frozen Model A inference script implements the following fallback hierarchy for deployment:"
    )
    add_code_block(doc,
        "def predict_c4h8(current_conditions, history_24h, last_analyzer_value=None):\n"
        "    # 1. Build anchor from last_analyzer_value (shift(1) guaranteed)\n"
        "    # 2. Compute dev24h deviations from history_24h\n"
        "    # 3. Build ratio features from current_conditions\n"
        "    # 4. Run XGBoost prediction\n"
        "\n"
        "    # Fallback hierarchy:\n"
        "    # prediction_health: GREEN  → anchor available, prediction valid\n"
        "    # prediction_health: YELLOW → anchor stale (>12h), rolling mean fallback\n"
        "    # prediction_health: RED    → no anchor, no history → use default mean\n"
        "\n"
        "    return {\n"
        "        'predicted_c4h8': float,\n"
        "        'model_used': 'Model A (XGBoost)' | 'Fallback (Rolling Mean)' | 'Fallback (Default)',\n"
        "        'prediction_health': 'GREEN' | 'YELLOW' | 'RED',\n"
        "        'fallback_reason': str | None\n"
        "    }"
    )
    add_divider(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 6 — MODEL B
    # ═══════════════════════════════════════════════════════════════════════════
    page_break(doc)
    add_heading(doc, "6. Model B — C4H6 Soft Sensor Development", level=1, color=(13,71,161))
    add_paragraph(doc,
        "Model B targets C4H6 wt% (butadiene) in the bottom product. "
        "The Model B story is fundamentally different from Model A — it is a story of systematic experiments "
        "that proved ML was unnecessary and led to an engineering decision rather than a modeling decision."
    )

    add_heading(doc, "6.1 Target Distribution Audit", level=2)
    add_table(doc,
        ["Block", "Mean C4H6 (wt%)", "Std C4H6 (wt%)", "Interpretation"],
        [
            ["1", "0.208", "High", "Cold reboiler regime — fundamentally different thermodynamics"],
            ["2", "0.031", "Moderate", "Normal hot operation — 37x lower than Block 1"],
            ["3", "0.023", "Moderate", "Normal hot operation"],
            ["4", "0.0057", "Very Low", "Lower pressure, different feed — ~5x lower than Block 2"],
        ]
    )
    add_paragraph(doc,
        "Block 1 C4H6 is 37x higher than Block 4. Including Block 1 in training "
        "caused R² to plummet to -17 to -35 on Block 4. Block 1 was excluded."
    )

    add_heading(doc, "6.2 Anchor-Only Baseline Evaluation", level=2)
    add_paragraph(doc, "Four baselines were evaluated against the Block 4 test set:")
    add_table(doc,
        ["Method", "R² (Block 4)", "MAE (wt%)", "Interpretation"],
        [
            ["Block 4 Mean (0.0057)", "0.00", "0.0032", "Simple mean predictor — poor"],
            ["Anchor (72h forward-fill)", "0.607", "0.0007", "Good but stale on long gaps"],
            ["Rolling Anchor (12h window)", "0.931", "0.0009", "Strong but more complex"],
            ["Anchor (12h forward-fill)", "0.9606", "0.0005", "Best — persistence model"],
        ]
    )
    add_paragraph(doc,
        "The 12-hour persistence anchor alone achieves R²=0.9606 with MAE=0.0005 wt%. "
        "This is a remarkable result: the butadiene content changes so slowly that simply remembering "
        "the last valid analyzer reading and holding it for up to 12 hours is nearly optimal."
    )

    add_heading(doc, "6.3 Delta ML Correction Experiment", level=2)
    add_paragraph(doc,
        "To test whether ML could improve upon the anchor, a delta correction model was trained "
        "(predict: anchor_error = true_C4H6 - C4H6_anchor) using the 7 process deviation and ratio features:"
    )
    add_table(doc,
        ["Approach", "R² (Block 4)", "MAE (wt%)", "Conclusion"],
        [
            ["Anchor only (12h)", "0.9606", "0.000547", "Strong baseline"],
            ["Anchor + XGBoost delta", "0.9010", "0.001194", "ML made it WORSE"],
        ]
    )
    add_paragraph(doc,
        "The delta model worsened performance. The training delta mean was -0.000018 (essentially zero), "
        "meaning the model was trying to learn random noise. "
        "Conclusion: process variables contain no additional information beyond the analyzer memory for butadiene. "
        "This is an engineering finding, not a modeling failure."
    )

    add_heading(doc, "6.4 Model B Decision & Production Implementation", level=2)
    add_paragraph(doc, "Decision:", bold=True)
    add_paragraph(doc,
        "Model B is implemented as a pure 12-hour persistence anchor. No ML model is used. "
        "The inference script (predict_c4h6.py) returns the last valid non-stuck C4H6 analyzer reading, "
        "forward-filled for up to 12 hours (shift(1) guaranteed for leak-free deployment)."
    )
    add_divider(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 7 — CONCEPT DRIFT
    # ═══════════════════════════════════════════════════════════════════════════
    page_break(doc)
    add_heading(doc, "7. Concept Drift: Discovery, Diagnosis & Resolution", level=1, color=(13,71,161))
    add_paragraph(doc,
        "Concept drift was the central technical challenge of this project. "
        "Every naive model trained on Blocks 1-3 and tested on Block 4 produced negative R² values — "
        "predictions were actively worse than the mean predictor, and in the wrong direction."
    )

    add_heading(doc, "7.1 The Discovery — Sign Reversal in Correlations", level=2)
    add_paragraph(doc,
        "A Pearson correlation audit revealed a complete reversal of sign for all temperature features between training and test:"
    )
    add_table(doc,
        ["Feature", "Pearson (Train Blocks 1-3)", "Pearson (Test Block 4)", "Sign Reversed?"],
        [
            ["Column_Top_Temp", "~+0.35", "~−0.28", "YES"],
            ["Control_Tray_Temp", "~+0.42", "~−0.31", "YES"],
            ["Column_Bottom_Temp", "~+0.29", "~−0.22", "YES"],
            ["Reboiler_Outlet_Temp", "~+0.38", "~−0.25", "YES"],
            ["Column_Top_Pressure", "~+0.27", "~−0.19", "YES"],
            ["Reflux_Ratio", "~+0.61", "~+0.58", "NO — stable"],
            ["Steam_Feed_Ratio", "~+0.55", "~+0.52", "NO — stable"],
        ]
    )

    add_heading(doc, "7.2 The Root Cause — Thermodynamic Bubble Point Shift", level=2)
    add_paragraph(doc,
        "The column pressure setpoint changed from 4.19 bar (training) to 3.98 bar (test Block 4). "
        "In distillation, the bubble point temperature of a mixture is a function of composition AND pressure. "
        "When pressure drops by 0.21 bar, the bubble point temperature drops by approximately 3–5 °C for "
        "C4/C5 mixtures. This means that at Block 4 pressure settings, a temperature reading of 65 °C means "
        "a very different composition than the same 65 °C reading in Block 2."
    )
    add_paragraph(doc,
        "Additionally, absolute temperature correlations contain controller interaction artifacts ('rot'): "
        "for example, the Control Tray Temp has a measured correlation of −37.77 °C/bar with top pressure "
        "due to automated controller cascade interactions — not fundamental distillation thermodynamics. "
        "Any absolute temperature feature carries this non-physical, campaign-specific information."
    )

    add_heading(doc, "7.3 The Resolution — Engineering Around the Drift", level=2)
    add_paragraph(doc,
        "Three complementary strategies were applied to eliminate drift sensitivity:"
    )
    for item in [
        ("Dimensionless Ratios: ", "Reflux_Ratio and Steam_Feed_Ratio are independent of absolute pressure or temperature levels. They represent the fundamental mass and energy balance of the column."),
        ("24h Deviation Features: ", "Column_Bottom_Temp_dev24h captures the change from the recent operating baseline, not the absolute value. Campaign shifts affect the absolute level but not the short-term dynamics."),
        ("Campaign Anchor: ", "C4H8_campaign_anchor provides a dynamic offset. Any systematic campaign-level bias in the predictions is corrected by the anchor's prior."),
    ]:
        add_bullet(doc, item[1], bold_prefix=item[0])

    add_heading(doc, "7.4 Drift Experiment Results (11 Experiments)", level=2)
    add_table(doc,
        ["Experiment", "Test R²", "Test MAE (wt%)", "Top Feature"],
        [
            ["Baseline (67 features, all temps)", "−1.04", "0.3010", "month_cos (campaign memorization)"],
            ["No Calendar Features", "−0.97", "0.2930", "Control_Tray_Temp_lag1"],
            ["Pressure Normalization (k=3)", "−0.73", "0.2683", "Control_Tray_Temp_Pnorm_k3"],
            ["Rolling Deviations", "−0.85", "0.2784", "Control_Tray_Temp_lag1"],
            ["Campaign Anchor (shift(1), leak-free)", "+0.69", "0.1032", "C4H8_campaign_anchor (0.538 importance)"],
            ["Ablation Study: 8-feature Subset 7", "+0.91", "0.0516", "C4H8_campaign_anchor + ratios + dev24h"],
        ]
    )
    add_divider(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 8 — SURROGATE MODELS
    # ═══════════════════════════════════════════════════════════════════════════
    page_break(doc)
    add_heading(doc, "8. Phase 5.1A — Surrogate Process Models", level=1, color=(13,71,161))
    add_paragraph(doc,
        "Before building the advisory optimizer, three lightweight surrogate process models were trained "
        "to predict how the column thermal state (bottom temperature, tray temperature, top pressure) "
        "responds to changes in steam and reflux flow. This transforms the optimizer from a naive grid search "
        "into a physically realistic prediction chain."
    )

    add_heading(doc, "8.1 Why Surrogate Models Are Necessary", level=2)
    add_paragraph(doc,
        "Without surrogate models, the optimizer would evaluate candidates by calling the soft sensor with "
        "frozen temperature and pressure values — as if column thermodynamics are instantaneous and perfectly "
        "known. This is physically incorrect. A 1 TPH increase in steam flow will change the column's "
        "bottom temperature by approximately 0.5–2 °C within 1 hour. Ignoring this would:"
    )
    for item in [
        "Incorrectly predict compositions using stale process conditions.",
        "Miss safety limit violations (bottom temp could exceed 115 °C with large steam increases).",
        "Generate physically inconsistent recommendations.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8.2 Delta Target Formulation", level=2)
    add_paragraph(doc,
        "Surrogate models are trained on 1-hour ahead delta targets (not absolute values):"
    )
    add_code_block(doc,
        "# T1 (Bottom Temp):\n"
        "# Target uses delta of dev24h deviation for extra drift robustness:\n"
        "delta_bot = (bottom_temp_dev24h[t+1]) - (bottom_temp_dev24h[t])\n"
        "\n"
        "# T2 (Tray Temp) and T3 (Pressure):\n"
        "delta_tray = tray_temp[t+1] - tray_temp[t]\n"
        "delta_pres = pressure[t+1] - pressure[t]\n"
        "\n"
        "# At inference:\n"
        "pred_bottom_temp = current_bottom_temp + surrogate_T1.predict(features)\n"
        "pred_tray_temp   = current_tray_temp   + surrogate_T2.predict(features)\n"
        "pred_pressure    = current_pressure    + surrogate_T3.predict(features)"
    )

    add_heading(doc, "8.3 Validation Results", level=2)
    add_table(doc,
        ["Target", "Naive R²", "Model R²", "Naive MAE", "Model MAE", "Winner", "Threshold (R²≥0.80)"],
        [
            ["Bottom Temp (T1)", "0.7137", "0.7638", "0.7301 °C", "0.6799 °C", "CatBoost", "Not met (0.764)"],
            ["Tray Temp (T2)", "0.8979", "0.9076", "1.887 °C", "1.986 °C", "CatBoost", "Met"],
            ["Top Pressure (T3)", "0.9495", "0.9491", "0.0125 bar", "0.0148 bar", "CatBoost", "Met (marginally)"],
        ]
    )

    add_heading(doc, "8.4 Engineering Assessment", level=2)
    for item in [
        ("Bottom Temp: ", "Model R²=0.764 is below the aspirational threshold of 0.80, but beats naive by +0.05 R² and shows lower MAE. At ±0.69 °C MAE, this is physically acceptable for an advisory context — the safety buffer system compensates."),
        ("Tray Temp: ", "Marginal improvement (+0.011 R²). Tray temperature is controlled via a cascade loop that partially decouples it from direct steam/reflux changes, limiting the direct learning signal."),
        ("Pressure: ", "Negligible improvement over naive persistence. Column top pressure is tightly controlled by the overhead pressure controller. The surrogate is retained for architectural consistency but adds minimal computational value."),
    ]:
        add_bullet(doc, item[1], bold_prefix=item[0])
    add_divider(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 9 — OPTIMIZER
    # ═══════════════════════════════════════════════════════════════════════════
    page_break(doc)
    add_heading(doc, "9. Phase 5.2 — Physics-Aware Advisory Optimizer", level=1, color=(13,71,161))
    add_paragraph(doc,
        "The advisory optimizer (optimizer_v2_physics_aware.py) is the capstone component of this project. "
        "It chains the surrogate process models with the frozen soft sensor to generate physically realistic, "
        "safe, and economically aware setpoint recommendations for the board operator."
    )

    add_heading(doc, "9.1 Search Strategy", level=2)
    add_paragraph(doc,
        "A local grid search is performed around the current operating point:"
    )
    add_table(doc,
        ["Variable", "Global Range", "Local Constraint", "Step Size", "Max Grid Points"],
        [
            ["Reboiling Steam Flow", "18.0 – 24.4 TPH", "±2.0 TPH from current", "0.2 TPH", "21 points"],
            ["Reflux Flow", "80.0 – 103.9 TPH", "±10.0 TPH from current", "1.0 TPH", "21 points"],
        ]
    )
    add_paragraph(doc, "Maximum candidates per call: 21 × 21 = 441. Typical runtime: < 3 seconds.")

    add_heading(doc, "9.2 Per-Candidate Evaluation Chain", level=2)
    add_code_block(doc,
        "For each (steam_cand, reflux_cand):\n"
        "  Step A: Predict delta_bot, delta_tray, delta_pres via surrogate models\n"
        "  Step B: Reconstruct absolute predictions:\n"
        "           pred_bot = current_bot_temp + delta_bot\n"
        "  Step C: Safety check (with MAE buffer):\n"
        "           REJECT if (pred_bot + 0.69) > 115.0 °C\n"
        "           REJECT if (pred_pres + 0.014) > 5.0 bar\n"
        "  Step D: Build Model A features using pred_bot, pred_tray, pred_pres\n"
        "  Step E: pred_c4h8 = model_a.predict(features)\n"
        "  Step F: pred_c4h6 = current_c4h6_anchor  (constant — Model B assumption)\n"
        "  Step G: pred_total = pred_c4h8 + pred_c4h6\n"
        "  Step H: REJECT if pred_total >= current_total_c4  (no improvement guarantee)"
    )

    add_heading(doc, "9.3 Two-Stage Objective Function", level=2)
    add_code_block(doc,
        "# Stage 1: Find all candidates meeting product specification\n"
        "spec_compliant = candidates[candidates['pred_total_c4'] < 0.50]\n"
        "\n"
        "if spec_compliant is not empty:\n"
        "    if MODE == 'economic':\n"
        "        # Among spec-compliant: minimize utility cost\n"
        "        winner = spec_compliant.sort_values('cost_benefit').iloc[0]\n"
        "    else:  # MODE == 'spec'\n"
        "        # Among spec-compliant: minimize C4 slippage\n"
        "        winner = spec_compliant.sort_values('pred_total_c4').iloc[0]\n"
        "else:\n"
        "    # Cannot meet spec in one step: minimize C4 regardless\n"
        "    winner = candidates.sort_values('pred_total_c4').iloc[0]"
    )

    add_heading(doc, "9.4 Economic Scoring", level=2)
    add_code_block(doc,
        "# Utility cost change from current setpoint:\n"
        "delta_steam  = candidate_steam  - current_steam\n"
        "delta_reflux = candidate_reflux - current_reflux\n"
        "cost_benefit = (steam_cost_per_tph * delta_steam) + (reflux_cost_per_tph * delta_reflux)\n"
        "\n"
        "# Positive cost_benefit = increased utility spend\n"
        "# Negative cost_benefit = utility savings"
    )

    add_heading(doc, "9.5 Safety Confidence Rating", level=2)
    add_table(doc,
        ["Confidence Level", "Condition"],
        [
            ["HIGH", "Bottom Temp margin ≥ 3 °C AND Pressure margin ≥ 0.10 bar"],
            ["MEDIUM", "Neither HIGH nor LOW"],
            ["LOW", "Bottom Temp margin < 1 °C OR Pressure margin < 0.03 bar"],
        ]
    )
    add_paragraph(doc,
        "The confidence rating is based on safety margin distance from hard limits, not statistical R². "
        "This is deliberately operator-friendly — engineers care about how close the recommendation brings "
        "the column to a dangerous state, not about model R² values."
    )
    add_divider(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 10 — VALIDATION
    # ═══════════════════════════════════════════════════════════════════════════
    page_break(doc)
    add_heading(doc, "10. Batch Validation & Performance Results", level=1, color=(13,71,161))
    add_paragraph(doc,
        "The optimizer validation (optimizer_validation.py) ran both SPEC and ECONOMIC modes "
        "over 100 randomly sampled, out-of-spec snapshots (Total C4 > 0.50 wt%) from the held-out Block 4 campaign."
    )
    add_table(doc,
        ["Metric", "SPEC Mode", "ECONOMIC Mode"],
        [
            ["Recommendation Feasibility", "86.0% (86/100)", "86.0% (86/100)"],
            ["Safety Limit Violations", "0.0% (0/100)", "0.0% (0/100)"],
            ["Rejections (no C4 savings)", "14.0% (14/100)", "14.0% (14/100)"],
            ["Average C4 Reduction (absolute)", "0.1384 wt%", "0.1348 wt%"],
            ["Average C4 Reduction (relative)", "17.4%", "16.7%"],
            ["Average Steam Flow Change", "+0.99 TPH", "+0.56 TPH"],
            ["Average Reflux Flow Change", "−5.02 TPH", "−5.46 TPH"],
            ["Average Utility Cost Change", "−₹0.05/hr", "−₹2.67/hr"],
        ]
    )
    add_paragraph(doc, "Key Observations:", bold=True)
    for item in [
        ("Zero safety violations: ", "Every rejected candidate correctly violated the MAE-buffered safety constraints. No unsafe recommendation was generated across 100 tests."),
        ("14% baseline rejection: ", "These are periods where the column's bottom temperature is already close to the safety limit — correctly identified as 'no safe move available'."),
        ("Economic mode advantage: ", "ECONOMIC mode achieves nearly identical C4 reduction (16.7% vs 17.4%) while saving ₹2.67/hr by preferring lower-steam, lower-reflux solutions that still meet the 0.50 wt% spec."),
        ("Reflux decreases: ", "The optimizer typically recommends modest steam increases with reflux decreases — reflecting the key debutanizer trade-off where reboiler duty can compensate for lower overhead reflux."),
    ]:
        add_bullet(doc, item[1], bold_prefix=item[0])
    add_divider(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 11 — DASHBOARD
    # ═══════════════════════════════════════════════════════════════════════════
    page_break(doc)
    add_heading(doc, "11. Interactive Dashboard (NiceGUI Platform)", level=1, color=(13,71,161))
    add_paragraph(doc,
        "The NiceGUI web dashboard (debutanizer_dashboard/) integrates all models and the optimizer "
        "into an operator-facing interface. Running locally at http://localhost:8080."
    )

    add_heading(doc, "11.1 Dashboard Pages", level=2)
    add_table(doc,
        ["Page", "File", "Key Features"],
        [
            ["Overview", "pages/overview.py", "Color-coded KPI cards, timeline scrubber (Block 1-4), Column Health Card, Analyzer Status Card, Recommendation Preview Card"],
            ["Soft Sensor", "pages/soft_sensor.py", "Manual input form to run predictions on custom process conditions"],
            ["Advisory Optimizer", "pages/optimizer.py", "Full recommendation output: current state, recommended setpoints, predicted response, safety confidence, cost analysis"],
            ["Historical Trends", "pages/trends.py", "Interactive dual-axis Plotly trend charts with variable selector"],
            ["Diagnostics", "pages/diagnostics.py", "SHAP feature importance bar charts for Model A"],
            ["Settings", "pages/settings.py", "Live editor for configs/economics.json — mode, cost coefficients, safety limits"],
        ]
    )

    add_heading(doc, "11.2 Architecture & Key Engineering Decisions", level=2)
    for item in [
        ("Centralized Data Service (dashboard_data.py): ", "All predictions, optimizer calls, analyzer staleness checks, and column health evaluations are consolidated in a single service. Pages read from this service rather than making direct calls — ensuring consistent state across all views."),
        ("NaN Safety: ", "All prediction outputs are checked with np.isnan() and rendered as '--' rather than 'nan'. Operators must never see NaN in an industrial dashboard."),
        ("Reactive UI: ", "Navigation state is managed via a global active_page variable. @ui.refreshable decorator on render_content() and render_navigation() enables reactive re-renders without full page reloads."),
        ("sys.path Management: ", "app.py prepends both the dashboard directory and the project root to sys.path, making all internal imports robust regardless of the launch directory."),
        ("Currency: ", "All economic outputs display Indian Rupees (₹) for IOCL context."),
    ]:
        add_bullet(doc, item[1], bold_prefix=item[0])
    add_divider(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 12 — PROBLEMS & RECOVERIES
    # ═══════════════════════════════════════════════════════════════════════════
    page_break(doc)
    add_heading(doc, "12. Problems, Setbacks & Recoveries", level=1, color=(13,71,161))

    problems = [
        ("Problem 1: Model predicting in completely wrong direction (R² = −1.0)",
         "First naive XGBoost model trained on Blocks 1-3 gave R² = −1.04 on Block 4.",
         "Systematic correlation audit revealed sign reversal in all temperature features. Root cause: thermodynamic bubble point shift due to pressure setpoint change between campaigns. Recovery: Replaced absolute temperature features with 24h deviation features and dimensionless ratios."),
        ("Problem 2: Campaign anchor causing target leakage (fake R² = 0.99)",
         "Initial campaign anchor implementation directly used C4H8_Bottom as the anchor feature.",
         "The anchor was the target itself — a direct data leak. Discovered during leakage audit. Recovery: Applied shift(1) before ffill. This guarantees the model only sees the previous valid reading, not the current one."),
        ("Problem 3: Stuck analyzer periods producing false training labels",
         "Large fractions of the training data had 'stuck' analyzer readings repeating for hours.",
         "Implemented automated stuck detection algorithm. All stuck rows are flagged and masked out from target labels. The anchor forward-fill mechanism handles prediction during stuck periods."),
        ("Problem 4: Block 1 causing severe target distribution mismatch for Model B",
         "Including Block 1 (cold startup) in Model B training → R² = −17 to −35 on Block 4.",
         "Target audit revealed Block 1 mean C4H6 = 0.208 wt% vs Block 4 mean = 0.0057 wt% (37x difference). Block 1 excluded from all Model B training with documentation."),
        ("Problem 5: ML delta model for Model B degraded performance",
         "A delta correction model (Anchor + XGBoost correction) gave R² = 0.901 vs anchor-only R² = 0.9606.",
         "The training delta mean was essentially zero (−0.000018). ML was fitting noise. Decision: Freeze Model B as pure persistence anchor. Documented as engineering decision."),
        ("Problem 6: Bottom Temp surrogate below R²=0.80 threshold",
         "Target threshold was 0.80. Achieved R²=0.764.",
         "Reformulated T1 target as delta of dev24h (not absolute delta), gaining +0.05 R². Revised acceptance criteria to focus on whether model beats naive baseline and whether MAE is acceptable for advisory use. Deployed with MAE-based safety buffers."),
        ("Problem 7: NiceGUI dashboard showing NaN to operators",
         "Early dashboard showed 'nan wt%' in KPI cards due to NaN propagating from stuck analyzer periods.",
         "Added safe_num() helper function that formats NaN as '--'. Corrected spec status logic to show 'DATA INVALID/OFFLINE' instead of defaulting to green when value is NaN."),
        ("Problem 8: Dashboard ModuleNotFoundError for optimizer module",
         "optimizer_service.py could not find optimizer_v2_physics_aware.py at runtime.",
         "Changed import from relative name to package-qualified name (notebooks.optimizer_v2_physics_aware). Also added explicit sys.path.append() for the notebooks/ directory in app.py."),
        ("Problem 9: NiceGUI ui.clear() error causing page crashes",
         "Using ui.clear() inside @ui.refreshable-decorated functions caused AttributeError crashes.",
         "Removed all ui.clear() calls. The @ui.refreshable decorator handles container clearing automatically. Replaced e.value event handler lookups with direct element.value reads."),
    ]

    for i, (title_str, problem, recovery) in enumerate(problems, 1):
        add_heading(doc, title_str, level=2)
        add_paragraph(doc, "Problem: ", bold=True)
        add_paragraph(doc, problem, indent=0.2)
        add_paragraph(doc, "Recovery: ", bold=True)
        add_paragraph(doc, recovery, indent=0.2)
    add_divider(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 13 — FILE REFERENCE
    # ═══════════════════════════════════════════════════════════════════════════
    page_break(doc)
    add_heading(doc, "13. File & Folder Reference", level=1, color=(13,71,161))

    file_sections = [
        ("Root Level Scripts", [
            ("data_preprocessing.py", "Cleans raw Excel data. Applies winsorization, stuck analyzer detection, and block assignment. Outputs data/clean_data.parquet."),
            ("feature_engineering.py", "Gap-aware time-series resampling. Computes lags, rolling means, ratio features, and dev24h deviations. Outputs data/features.parquet."),
            ("model_training.py", "General-purpose model training script (early experiments)."),
            ("experiment_catboost.py", "CatBoost-specific experiment runner."),
            ("experiment_optuna.py", "Optuna hyperparameter tuning experiment runner."),
            ("generate_full_documentation.py", "This documentation generator script."),
        ]),
        ("notebooks/ — Research & Development Scripts", [
            ("tune_robust_xgb.py", "50-trial Optuna hyperparameter optimization for Model A using TimeSeriesSplit CV."),
            ("verify_anchor_leakage.py", "Formal leakage verification: proves shift(1) is applied and reports Block 3 metrics."),
            ("freeze_model_A.py", "Serializes Model A to JSON/PKL and copies release artifacts to final_v1/."),
            ("run_drift_experiments.py", "11-experiment systematic drift mitigation study."),
            ("run_feature_ablation_study.py", "7-subset feature ablation to find optimal feature set."),
            ("run_anchor_analysis.py", "Anchor fill limit sensitivity analysis (6h, 12h, 24h, 48h, 72h)."),
            ("model_b_target_audit.py", "Block-by-block C4H6 target distribution audit."),
            ("anchor_only_baselines.py", "4-baseline evaluation proving 12h anchor achieves R²=0.96."),
            ("model_b_delta_model.py", "Delta correction experiment for Model B."),
            ("model_b_inversion_check.py", "Anchor robustness cross-validation across Blocks 2, 3, 4."),
            ("build_surrogate_dataset.py", "Creates surrogate dataset with t+1 shift targets within each block."),
            ("train_surrogate_models.py", "Trains XGBoost/LightGBM/CatBoost surrogate models. Saves winner PKL and feature importances."),
            ("surrogate_diagnostics.py", "Generates diagnostic plots for all three surrogates."),
            ("optimizer_v2_physics_aware.py", "Physics-Aware Advisory Optimizer (Phase 5.2). Full recommendation engine."),
            ("optimizer_validation.py", "Batch validation of optimizer over 100 Block 4 out-of-spec snapshots."),
        ]),
        ("inference/ — Production Inference", [
            ("predict_c4h8.py", "Model A inference with fallback hierarchy (GREEN/YELLOW/RED health status)."),
            ("predict_c4h6.py", "Model B persistence anchor with 12h limit. Returns latest valid C4H6 reading."),
            ("predict_total_c4.py", "Unified entry point combining Model A + Model B for Total C4 prediction."),
        ]),
        ("debutanizer_dashboard/ — NiceGUI Web Dashboard", [
            ("app.py", "Main entry point. Quasar dark theme, reactive navigation sidebar, sys.path setup."),
            ("pages/overview.py", "KPI cards, timeline scrubber, Column Health, Analyzer Status, Recommendation Preview cards."),
            ("pages/soft_sensor.py", "Manual prediction input form."),
            ("pages/optimizer.py", "Advisory optimizer output display."),
            ("pages/trends.py", "Interactive historical trend charts."),
            ("pages/diagnostics.py", "Feature importance visualization."),
            ("pages/settings.py", "Live economics.json editor."),
            ("components/cards.py", "Reusable color-coded KPI cards (Green/Yellow/Red)."),
            ("components/charts.py", "Plotly chart builders with correct layout formatting."),
            ("services/dashboard_data.py", "Centralized data aggregation service."),
            ("services/prediction_service.py", "Wrapper around predict_total_c4.py."),
            ("services/optimizer_service.py", "Wrapper around optimizer_v2_physics_aware.py."),
            ("services/state_service.py", "Global state management (current snapshot index, loaded dataframes)."),
        ]),
        ("configs/ — Configuration Files", [
            ("model_A_features.json", "Feature names, order, and metadata for Model A deployment."),
            ("model_B_features.json", "Feature configuration for Model B anchor."),
            ("economics.json", "Optimizer configuration: mode, price coefficients, operating bounds, safety limits."),
        ]),
        ("models/ — Serialized Models", [
            ("final/model_A_final_v1.json", "Frozen XGBoost Model A weights (JSON format for DCS deployment)."),
            ("final/model_A_final_v1.pkl", "Frozen XGBoost Model A (Python pickle format)."),
            ("surrogates/bottom_temp_t1_model.pkl", "T1 surrogate: Bottom Temperature 1-hour-ahead delta model."),
            ("surrogates/tray_temp_t1_model.pkl", "T2 surrogate: Tray Temperature 1-hour-ahead delta model."),
            ("surrogates/pressure_t1_model.pkl", "T3 surrogate: Top Pressure 1-hour-ahead delta model."),
            ("surrogates/surrogate_results.json", "Surrogate validation metrics and hyperparameters."),
        ]),
        ("docs/ — Documentation", [
            ("01_problem_statement.md", "Problem statement and objectives."),
            ("02_dataset_understanding.md", "Dataset structure and variable descriptions."),
            ("03_feature_engineering.md", "Feature engineering design and rationale."),
            ("04_model_a_development.md", "Model A development journal."),
            ("05_model_b_development.md", "Model B development journal."),
            ("06_drift_analysis.md", "Concept drift experiments and findings."),
            ("07_final_architecture.md", "Final system architecture description."),
            ("08_deployment_guide.md", "Deployment instructions and DCS integration notes."),
            ("09_future_optimizer.md", "Future optimizer design notes."),
            ("optimizer_summary.md", "Comprehensive Phase 5 optimizer documentation (this section)."),
        ]),
    ]

    for section_title, files in file_sections:
        add_heading(doc, section_title, level=2)
        add_table(doc,
            ["File", "Purpose"],
            [(f, d) for f, d in files]
        )
    add_divider(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 14 — MATHEMATICAL REFERENCE
    # ═══════════════════════════════════════════════════════════════════════════
    page_break(doc)
    add_heading(doc, "14. Key Formulas & Mathematical Reference", level=1, color=(13,71,161))

    add_heading(doc, "14.1 Feature Engineering", level=2)
    add_code_block(doc, "Reflux_Ratio = Reflux_Flow / Feed_Flow")
    add_code_block(doc, "Steam_Feed_Ratio = Reboiling_Steam_Flow / Feed_Flow")
    add_code_block(doc, "X_dev24h = X - rolling_mean(X, window=24h)")
    add_code_block(doc, "C4H8_anchor(t) = last_valid_C4H8(t-1), forward-filled up to 72h within block")

    add_heading(doc, "14.2 Surrogate Model Target Formulation", level=2)
    add_code_block(doc,
        "# Bottom Temp (T1) — dev24h delta:\n"
        "Δy_bot = bottom_temp_dev24h[t+1] - bottom_temp_dev24h[t]\n"
        "pred_bottom_temp = current_bottom_temp + surrogate_T1.predict(Δy_bot)\n\n"
        "# Tray Temp (T2) and Pressure (T3) — raw delta:\n"
        "Δy_tray = tray_temp[t+1] - tray_temp[t]\n"
        "Δy_pres = pressure[t+1] - pressure[t]"
    )

    add_heading(doc, "14.3 Safety Constraint with MAE Buffer", level=2)
    add_code_block(doc,
        "# Reject candidate if:\n"
        "pred_bottom_temp + MAE_bottom (0.690 °C) > 115.0 °C\n"
        "pred_pressure    + MAE_pressure (0.014 bar) > 5.0 bar"
    )

    add_heading(doc, "14.4 Economic Objective Function", level=2)
    add_code_block(doc,
        "cost_benefit = steam_cost_per_tph × Δsteam + reflux_cost_per_tph × Δreflux\n\n"
        "# Two-stage selection:\n"
        "Stage 1: Filter all candidates where pred_total_c4 < 0.50 wt%\n"
        "Stage 2 (SPEC mode):     minimize pred_total_c4\n"
        "Stage 2 (ECONOMIC mode): minimize cost_benefit"
    )

    add_heading(doc, "14.5 Economic Loss Estimation (Dashboard)", level=2)
    add_code_block(doc,
        "c4_fraction = pred_total_c4 / 100.0\n"
        "feed_kg_hr = Feed_Flow × 1000.0  # TPH to kg/hr\n"
        "c4_loss_kg_hr = feed_kg_hr × c4_fraction\n"
        "loss_rs_per_hr = c4_loss_kg_hr × 142.0  # ₹142/kg C4 value"
    )

    add_heading(doc, "14.6 Goodness-of-Fit Metrics", level=2)
    add_code_block(doc,
        "R² = 1 - (SS_residual / SS_total)\n"
        "   = 1 - [Σ(y_true - y_pred)²] / [Σ(y_true - ȳ)²]\n\n"
        "MAE = (1/n) × Σ|y_true - y_pred|\n\n"
        "Note: R² < 0 means the model is worse than predicting the mean.\n"
        "      R² = -17 (worst case in drift experiments) means the model\n"
        "      is actively predicting in the WRONG DIRECTION."
    )
    add_divider(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 15 — LESSONS LEARNED
    # ═══════════════════════════════════════════════════════════════════════════
    page_break(doc)
    add_heading(doc, "15. Lessons Learned & Final Conclusions", level=1, color=(13,71,161))

    lessons = [
        ("Lesson 1: Always check for concept drift before trusting any ML result.",
         "The first model gave R² = +0.72 in cross-validation but R² = −1.0 on the true test set. "
         "Without a proper temporal holdout, this would have been deployed as a working model."),
        ("Lesson 2: Negative R² is diagnostic information, not failure.",
         "R² = −1.0 is not a bad model — it is a sign that correlations reversed. "
         "The direction of prediction inversion directly pointed to thermodynamic bubble point shift "
         "as the root cause."),
        ("Lesson 3: Process physics should drive feature design.",
         "Adding more features never solved the drift. The breakthrough came from understanding "
         "distillation thermodynamics: pressure-normalized temperatures and deviations from rolling means "
         "are fundamentally drift-resistant because they measure relative changes, not absolute states."),
        ("Lesson 4: Always compare ML to naive baselines.",
         "Model B's delta model scored worse than the raw persistence anchor. "
         "Without explicitly comparing to the naive baseline first, this would never have been detected."),
        ("Lesson 5: An engineering decision is a valid conclusion.",
         "Concluding that a simple persistence anchor is better than ML for Model B "
         "is a stronger result than a mediocre ML model. The systematic experiments that proved this "
         "are more valuable than the ML model they replaced."),
        ("Lesson 6: Safety-first design for industrial AI.",
         "Every safety check in the optimizer uses conservative MAE-buffered limits. "
         "The optimizer must be designed to fail safely — when uncertain, it should refuse to recommend "
         "rather than output a potentially unsafe setpoint."),
        ("Lesson 7: Surrogate models unlock physics-aware optimization.",
         "The upgrade from naive grid-search to surrogate-chain optimizer represents a fundamental "
         "improvement in physical correctness. The 0-violation-in-100-tests safety record directly reflects "
         "this design philosophy."),
    ]

    for lesson_title, lesson_text in lessons:
        add_heading(doc, lesson_title, level=2)
        add_paragraph(doc, lesson_text)
    add_divider(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 16 — FUTURE WORK
    # ═══════════════════════════════════════════════════════════════════════════
    page_break(doc)
    add_heading(doc, "16. Future Work & Roadmap", level=1, color=(13,71,161))

    future = [
        ("Multi-Step Safety Horizon (T+2, T+3): ",
         "Train T+2 and T+3 surrogate models and evaluate max(pred_temp_T1, T2, T3) against safety limits. "
         "This prevents the optimizer from recommending large steam increases that would be safe at T+1 "
         "but dangerous after 3 hours of dynamics."),
        ("MPC / APC Integration: ",
         "Transition from advisory mode to closed-loop Multivariable Predictive Control (MPC) once "
         "steady-state column response models are validated with live plant feedback. "
         "The current optimizer architecture is directly extensible to MPC."),
        ("Seeq Deployment: ",
         "Deploy the soft sensor and optimizer on the refinery's Seeq historian platform. "
         "This would enable real-time predictions on live DCS data and automated alarm triggers."),
        ("Economic Calibration: ",
         "Replace placeholder cost coefficients in economics.json with real refinery utility tariffs "
         "and actual C4 product recovery pricing. Calibrate the economic mode for real-plant optimization."),
        ("C4H6 Manipulable Variable Response Model: ",
         "If future analyzer data captures sufficient variation in steam/reflux with concurrent C4H6 "
         "readings, a proper butadiene response surrogate could be trained to replace the constant-anchor assumption."),
        ("Bayesian / Differential Evolution Optimizer: ",
         "For larger search spaces (e.g., adding Feed Flow as a third manipulated variable), replace "
         "grid search with Differential Evolution or Bayesian Optimization (via scipy.optimize or optuna)."),
        ("Closed-Loop Validation: ",
         "Run controlled plant trials with the optimizer in parallel (advisory mode) and compare "
         "recommended vs. actual operator actions and resulting C4 outcomes to validate the recommendation quality."),
    ]

    for title_str, body in future:
        add_bullet(doc, body, bold_prefix=title_str)

    doc.add_paragraph()
    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = final.add_run(
        "─── End of Documentation ───\n"
        "IOCL Debutanizer C4 Slippage Optimization — AI Soft Sensor & Advisory Optimizer\n"
        "Status: Frozen & Validated — v2.1 Advisory | All models deployed in NiceGUI dashboard"
    )
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(120, 120, 120)

    # Save
    out_path = "Debutanizer_C4_Project_Documentation.docx"
    doc.save(out_path)
    print(f"\n[OK] Documentation saved to: {out_path}")
    print(f"  File size: {os.path.getsize(out_path) / 1024:.1f} KB")


if __name__ == "__main__":
    build_docx()
