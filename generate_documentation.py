"""
generate_documentation.py
=========================
Generates a comprehensive DOCX documentation file for the
Debutanizer C4 Slippage Optimization project.
"""

import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.dml.color import ColorFormat
import docx.opc.constants

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color:
        run = h.runs[0] if h.runs else h.add_run(text)
        run.font.color.rgb = RGBColor(*color)
    return h

def add_para(doc, text, bold=False, italic=False, size=11, color=None, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_numbered(doc, text, level=0):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_code_block(doc, code_text):
    """Add a code block with light gray background."""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    doc.add_paragraph()

def add_table(doc, headers, rows, bold_header=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        if bold_header:
            for run in hdr_cells[i].paragraphs[0].runs:
                run.bold = True
            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # shade header
            tc = hdr_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '2E4057')
            tcPr.append(shd)
    # Data rows
    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
    doc.add_paragraph()
    return table

def add_image(doc, img_path, caption='', width=6.0):
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(width))
        if caption:
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.italic = True
                run.font.size = Pt(9)
    else:
        p = doc.add_paragraph(f'[Image not found: {os.path.basename(img_path)}]')
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

def add_horizontal_line(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E4057')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_note_box(doc, text, note_type='NOTE'):
    colors = {
        'NOTE': (0x17, 0x67, 0xCC),
        'WARNING': (0xCC, 0x77, 0x00),
        'IMPORTANT': (0x8B, 0x00, 0x00),
        'TIP': (0x00, 0x80, 0x40),
    }
    color = colors.get(note_type, (0x33, 0x33, 0x33))
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f'{note_type}: {text}')
    run.bold = True
    run.font.color.rgb = RGBColor(*color)
    run.font.size = Pt(10)
    doc.add_paragraph()

# ============================================================
# DOCUMENT SECTIONS
# ============================================================

BASE_DIR = r'c:\Users\KIIT\OneDrive\Desktop\DEBUTANIZER-model'
IMG_DIR = os.path.join(BASE_DIR, 'experiments', 'diagnostics')
NB_DIR = os.path.join(BASE_DIR, 'notebooks')
FINAL_DIR = os.path.join(BASE_DIR, 'final_v1', 'models')


def build_cover(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading('Debutanizer Column', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x2E, 0x40, 0x57)
    title.runs[0].font.size = Pt(32)

    sub = doc.add_paragraph('AI-Based Soft Sensor for C4 Slippage Optimization')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(18)
    sub.runs[0].font.color.rgb = RGBColor(0x05, 0x96, 0x86)

    doc.add_paragraph()
    subtitle2 = doc.add_paragraph('Comprehensive Technical Documentation & Project Walkthrough')
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2.runs[0].italic = True
    subtitle2.runs[0].font.size = Pt(13)

    doc.add_paragraph()
    doc.add_paragraph()

    meta = [
        ('Project', 'IOCL Debutanizer Column Soft-Sensor'),
        ('Target', 'Total C4 Slippage (C4H8 + C4H6) Prediction'),
        ('Specification', '<= 0.50 wt% Total C4 in Bottom Product'),
        ('Model A Final R²', '0.9074 (XGBoost, Block 4 Test)'),
        ('Model A Final MAE', '0.0516 wt%'),
        ('Model B', 'Deterministic Anchor Tracker (R² = 0.9606)'),
        ('Version', 'final_v1 (Locked & Frozen)'),
        ('Status', 'Production Ready'),
    ]
    add_table(doc, ['Parameter', 'Value'], meta)
    doc.add_page_break()


def build_toc(doc):
    add_heading(doc, 'Table of Contents', 1)
    toc_items = [
        '1. Project Overview & Industrial Context',
        '2. Process Description: Debutanizer Column',
        '3. Dataset Overview',
        '4. Data Preprocessing (Phase 1)',
        '5. Feature Engineering (Phase 2)',
        '6. Initial Model Training & The Discovery of Failure (Phase 3)',
        '7. Root Cause Analysis: Concept Drift & Covariate Shift',
        '8. The Breakthrough: Campaign Anchor & Feature Ablation',
        '9. Model A (C4H8) — Final Development',
        '10. Model B (C4H6) — Development & Conclusion',
        '11. Final Architecture & Dual Soft-Sensor System',
        '12. Production Inference & Fallback Logic',
        '13. Deployment Guide',
        '14. Future Optimizer Integration',
        '15. File & Folder Structure Reference',
        '16. Key Metrics Summary',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0.5)
        p.runs[0].font.size = Pt(11)
    doc.add_page_break()


def section_01_overview(doc):
    add_heading(doc, '1. Project Overview & Industrial Context', 1)

    add_para(doc,
        'This project was commissioned by the Indian Oil Corporation Limited (IOCL) with the objective '
        'of deploying an AI-based virtual soft-sensor system for real-time prediction and minimization '
        'of C4 slippage in the Debutanizer Column bottom product.',
        size=11)
    doc.add_paragraph()

    add_heading(doc, '1.1 Formal Objective', 2)
    add_para(doc, 'Build a high-fidelity, real-time soft-sensor that predicts Total C4 composition '
             'in the column bottoms:', size=11)
    add_para(doc, '    Total C4 (wt%) = C4H8 (wt%) + C4H6 (wt%)', bold=True, size=11, indent=1)
    doc.add_paragraph()

    add_heading(doc, '1.2 Why This Matters — Industrial Stakes', 2)
    rows = [
        ('Product Off-Spec', 'Bottom product must be <= 0.50 wt% Total C4. 39.9% of historical readings exceeded this.'),
        ('Butadiene Poisoning', 'C4H6 at high levels poisons downstream polymerization catalysts.'),
        ('Steam Cost (Butene)', 'Over-fractionating to reduce C4H8 wastes reboiling steam (energy cost).'),
        ('Analyzer Delay', 'GC analyzer cycle: 12 min. Offline or stuck periods up to 37 days observed.'),
        ('Financial Impact', 'Losses of INR 1-10 Crore/hour possible during undetected slippage events.'),
    ]
    add_table(doc, ['Challenge', 'Detail'], rows)

    add_heading(doc, '1.3 Solution Architecture', 2)
    add_para(doc,
        'A dual-model virtual analyzer is deployed that runs every hour:', size=11)
    add_bullet(doc, 'Model A: XGBoost regressor (ML-based) for C4H8 (Butene) prediction.')
    add_bullet(doc, 'Model B: Deterministic analyzer-tracking state estimator for C4H6 (Butadiene).')
    add_bullet(doc, 'Combined output: Total C4 with GREEN/YELLOW/RED health flag for DCS dashboard.')
    doc.add_page_break()


def section_02_process(doc):
    add_heading(doc, '2. Process Description: Debutanizer Column', 1)

    add_heading(doc, '2.1 What a Debutanizer Does', 2)
    add_para(doc,
        'A Debutanizer is a distillation column designed to separate mixed C4 hydrocarbons '
        '(Butene C4H8, Butadiene C4H6, Butane, Isobutene) from heavier C5+ components '
        '(Pentane, Hexane, and heavier fractions). The goal is to keep C4 components in the '
        'overhead product, and minimize their "slip" into the bottom C5+ product stream.', size=11)
    doc.add_paragraph()

    add_heading(doc, '2.2 Key Process Variables (Raw Inputs)', 2)
    rows = [
        ('Feed_Flow', 'TPH', 'Column feed rate from upstream DP Bottom unit (Level-controlled)'),
        ('Reboiling_Steam_Flow', 'TPH', 'LP desuperheater steam providing reboiling duty (heat input)'),
        ('Reflux_Flow', 'TPH', 'Overhead condensate recycled back to column top (cooling input)'),
        ('Column_Bottom_Temp', 'deg C', 'Temperature at column bottom — indicator of stripping efficiency'),
        ('Control_Tray_Temp', 'deg C', 'Key cascade control temperature on intermediate tray'),
        ('Column_Top_Temp', 'deg C', 'Overhead temperature, indication of light component concentration'),
        ('Column_Top_Pressure', 'kg/cm2g', 'Operating pressure governing vapor-liquid equilibrium'),
        ('C4H8_Bottom', 'wt%', 'GC analyzer reading of Butene in column bottoms (TARGET A)'),
        ('C4H6_Bottom', 'wt%', 'GC analyzer reading of Butadiene in column bottoms (TARGET B)'),
    ]
    add_table(doc, ['Variable', 'Unit', 'Physical Role'], rows)

    add_heading(doc, '2.3 Process Physics', 2)
    add_para(doc, 'The core physical relationships governing the separation:', size=11)
    add_bullet(doc, 'HEAT INPUT (Steam): Higher reboiling steam raises bottom temp, driving C4 upward.')
    add_bullet(doc, 'REFLUX (Cooling): Higher reflux flow increases overhead condensation, improving separation.')
    add_bullet(doc, 'PRESSURE: Higher pressure raises all boiling points. C4H6 (Butadiene BP = -4.4C) is very volatile.')
    add_bullet(doc, 'PROCESS DELAY: Column requires 30-120 minutes to respond to control changes. Past conditions matter more than present.')
    doc.add_paragraph()

    add_heading(doc, '2.4 Operating Regimes Discovered', 2)
    rows = [
        ('Block 1 (2023)', '35.9 deg C (Cold)', 'Reboiler bypass / non-fractionating regime. C4H6 mean = 0.208 wt%'),
        ('Block 2 (2024)', '108.0 deg C (Hot)', 'Full fractionation, high steam. C4H6 mean = 0.031 wt%'),
        ('Block 3 (2024)', '93.2 deg C (Mixed)', 'Transition regime. C4H6 mean = 0.023 wt%'),
        ('Block 4 (2025-26)', '71.7 deg C (Mixed)', 'Production campaign. C4H6 mean = 0.0057 wt% (collapsed!)')
    ]
    add_table(doc, ['Campaign Block', 'Mean Reboiler Temp', 'Regime Description'], rows)

    add_note_box(doc,
        'Block 1 operated with a cold reboiler (mean 35.9 C vs 71-108 C in later blocks). '
        'This represents a completely different plant operating mode where fractionation was inactive. '
        'C4H6 mean (0.208 wt%) is 37x higher than Block 4 (0.0057 wt%). This is NOT drift - it is a different plant state.',
        'IMPORTANT')
    doc.add_page_break()


def section_03_dataset(doc):
    add_heading(doc, '3. Dataset Overview', 1)

    add_heading(doc, '3.1 Raw Data Structure', 2)
    add_para(doc,
        'The dataset was provided as an Excel file (9.DB DATA -B.xlsx) from the plant historian '
        '(Exaquantum) covering approximately 3 years of hourly process data.', size=11)

    rows = [
        ('Total Rows (after cleaning)', '11,343 hours'),
        ('Sampling Frequency', '1 hour (consistent)'),
        ('Date Range', 'April 2023 to April 2026'),
        ('Raw Columns', '11 (DateTime + 10 process variables)'),
        ('Total Gaps (Plant Shutdowns)', '2 major gaps (376 days and 258 days)'),
    ]
    add_table(doc, ['Property', 'Value'], rows)

    add_heading(doc, '3.2 Campaign Blocks & Gaps', 2)
    rows = [
        ('Block 1', '2023-04-16', '2023-08-31', '3,288', '137.0 days', '35.9 C (Cold)'),
        ('—', '376-day gap', '—', '—', '—', '—'),
        ('Block 2', '2024-09-11', '2024-10-11', '738', '30.7 days', '108.0 C (Hot)'),
        ('—', '43-hour gap', '—', '—', '—', '—'),
        ('Block 3', '2024-10-13', '2024-11-15', '803', '33.4 days', '93.2 C (Mixed)'),
        ('—', '258-day gap', '—', '—', '—', '—'),
        ('Block 4', '2025-08-01', '2026-04-30', '6,514', '272.0 days', '71.7 C (Mixed)'),
    ]
    add_table(doc, ['Block', 'Start', 'End', 'Rows', 'Duration', 'Reboiler Mean Temp'], rows)

    add_heading(doc, '3.3 Analyzer Quality Problems', 2)
    add_para(doc,
        'Both GC analyzers exhibited significant reliability issues, revealed through stuck-reading detection '
        '(sequences where identical values persisted for >= 12 consecutive hours):', size=11)
    rows = [
        ('C4H8_Bottom', '7.9%', 'Stuck at 0.034 wt% (frozen low) AND 1.262 wt% (frozen high)'),
        ('C4H6_Bottom', '33.0%', 'Stuck at exactly 0.000 wt% (sensor frozen). Real C4H6 is NEVER zero.'),
    ]
    add_table(doc, ['Analyzer', 'Stuck %', 'Pattern'], rows)

    add_heading(doc, '3.4 Analyzer Health Categories', 2)
    rows = [
        ('GOOD', '7,928', '69.9%', 'Both analyzers changed within 12h'),
        ('WARNING', '579', '5.1%', 'Unchanged for 12-24h'),
        ('BAD', '2,836', '25.0%', 'At least one analyzer flatlined > 24h'),
    ]
    add_table(doc, ['Status', 'Rows', 'Percentage', 'Meaning'], rows)
    add_note_box(doc, '25% of the entire dataset has at least one analyzer in BAD state. '
                 'Operators are "flying blind" a quarter of the time. This is a key value-add '
                 'for the soft-sensor dashboard.', 'IMPORTANT')

    add_heading(doc, '3.5 Target Distribution Statistics', 2)
    rows = [
        ('C4H8_Bottom (Train)', '4,353', '0.444', '0.032', '1.528', '0.310'),
        ('C4H8_Bottom (Test/Block4)', '6,093', '0.428', '0.072', '1.528', '0.276'),
        ('C4H6_Bottom (Train)', '3,577', '0.1395', '0.001', '0.669', '0.162'),
        ('C4H6_Bottom (Test/Block4)', '2,974', '0.0057', '0.001', '0.380', '0.010'),
    ]
    add_table(doc, ['Target', 'Count', 'Mean', 'Min', 'Max', 'Std Dev'], rows)
    add_note_box(doc,
        'C4H6 test mean is 24x lower than train mean (0.0057 vs 0.1395 wt%). '
        'This "target collapse" makes standard ML completely fail for C4H6 on Block 4. '
        'This became the central problem of Model B development.',
        'WARNING')
    doc.add_page_break()


def section_04_preprocessing(doc):
    add_heading(doc, '4. Data Preprocessing (Phase 1)', 1)

    add_heading(doc, '4.1 Objectives', 2)
    add_bullet(doc, 'Parse the raw Excel file with multi-row headers.')
    add_bullet(doc, 'Detect and remove plant shutdown periods (all-zero rows).')
    add_bullet(doc, 'Detect and flag stuck analyzer readings using floating-point-safe comparison.')
    add_bullet(doc, 'Apply IQR-based winsorization to clip extreme process values.')
    add_bullet(doc, 'Assign campaign block labels (Data_Block) to each row.')
    add_bullet(doc, 'Compute Analyzer_Health status (GOOD/WARNING/BAD).')
    add_bullet(doc, 'Save validated output to data/clean_data.parquet.')
    doc.add_paragraph()

    add_heading(doc, '4.2 Key Implementation: Floating-Point Stuck Detection', 2)
    add_para(doc,
        'The initial implementation used .diff().eq(0) to detect stuck readings. '
        'This was a critical bug because historian software writes values with tiny floating-point '
        'noise (e.g. 0.07980082929134369 vs 0.0798008292913436). np.isclose() was required instead:', size=11)
    add_code_block(doc, """
# WRONG (misses float-noise duplicates):
stuck_mask = series.diff().eq(0)

# CORRECT:
stuck_mask = np.isclose(series.diff().fillna(1.0), 0.0, atol=1e-6)
""")
    add_para(doc, 'Impact: Additional 15 stuck C4H6 rows caught (4,125 -> 4,140 total).', size=10, italic=True)

    add_heading(doc, '4.3 Shutdown Detection', 2)
    add_para(doc,
        '56 rows where ALL process variables were effectively zero (< epsilon = 0.01) '
        'were identified as plant shutdowns and removed. Using epsilon (not exact zero) '
        'future-proofs against historian export variations during plant ramp-down.', size=11)

    add_heading(doc, '4.4 Campaign Block Assignment', 2)
    add_para(doc,
        'Blocks were identified using datetime gaps > 24 hours between consecutive readings. '
        'Each row is labeled Data_Block = 1, 2, 3, or 4. This label is used throughout '
        'the project to prevent data leakage across shutdowns.', size=11)

    add_heading(doc, '4.5 Phase 1 Output', 2)
    rows = [
        ('Output File', 'data/clean_data.parquet'),
        ('Rows', '11,343'),
        ('Columns', '24 (original 11 + 13 derived metadata columns)'),
        ('Shutdown rows removed', '56'),
        ('C4H8 stuck rows flagged', '897'),
        ('C4H6 stuck rows flagged', '4,140'),
        ('Analyzer_Health: GOOD', '7,928 (69.9%)'),
        ('Analyzer_Health: BAD', '2,836 (25.0%)'),
    ]
    add_table(doc, ['Item', 'Value'], rows)
    doc.add_page_break()


def section_05_features(doc):
    add_heading(doc, '5. Feature Engineering (Phase 2)', 1)

    add_heading(doc, '5.1 Critical Rule: Block-Aware Lag Computation', 2)
    add_note_box(doc,
        'The SINGLE most important implementation rule: ALL lag and rolling computations '
        'must use groupby("Data_Block") to prevent information from crossing the 376-day '
        'shutdown gap. A global shift(1) would silently corrupt the first rows of each block.',
        'IMPORTANT')
    add_code_block(doc, """
# WRONG - leaks Block 1 data into Block 2 first row:
df[col].shift(1)

# CORRECT - block-isolated lags:
df.groupby("Data_Block")[col].shift(1)
""")

    add_heading(doc, '5.2 Mass & Energy Balance Ratios', 2)
    add_para(doc, 'Raw flow rates depend on column throughput. Dimensionless ratios '
             'remove this dependence and capture the physical intensity of separation:', size=11)
    rows = [
        ('Reflux_Ratio', 'Reflux_Flow / Feed_Flow', 'Separation cooling per unit feed. r = 0.285 with C4H8.'),
        ('Steam_Feed_Ratio', 'Reboiling_Steam_Flow / Feed_Flow', 'Separation heat input per unit feed. r = 0.324 with C4H8.'),
    ]
    add_table(doc, ['Feature', 'Formula', 'Significance'], rows)
    add_note_box(doc,
        'Key finding: Steam_Feed_Ratio (r=0.324) has 6x stronger correlation with C4H8 than '
        'raw steam flow alone (r=0.016). This directly validates the physics-based '
        'feature engineering approach.', 'TIP')

    add_heading(doc, '5.3 24-Hour Deviation Features (dev24h)', 2)
    add_para(doc,
        'Absolute temperatures, flows, and pressures drift over time due to seasonal '
        'variations, operator setpoint changes, or heat exchanger fouling. The deviation '
        'from a 24-hour rolling mean captures CHANGES rather than absolute state:', size=11)
    add_code_block(doc, """
Feature_dev24h = Feature_t - Rolling_Mean(Feature, 24 hours)

# Example: Steam deviation
Reboiling_Steam_Flow_dev24h = Reboiling_Steam_Flow[t] - mean(Reboiling_Steam_Flow[t-24:t])
""")
    add_para(doc, 'Variables with dev24h computed: Reboiling_Steam_Flow, Reflux_Flow, '
             'Column_Bottom_Temp, Control_Tray_Temp, Column_Top_Pressure.', size=10, italic=True)

    add_heading(doc, '5.4 Pressure-Normalized Temperatures (Pnorm)', 2)
    add_para(doc,
        'Boiling points shift with operating pressure (vapor-liquid equilibrium). '
        'A temperature reading means different things at different pressures. '
        'Thermodynamic correction normalizes temperatures to a reference pressure:', size=11)
    add_code_block(doc, """
Temp_Pnorm = Temp_t - (Column_Top_Pressure_t - P_ref) * k

where:
  P_ref = 4.05 kg/cm2g  (training-period mean pressure)
  k     = 3, 5, or 10 deg C / bar  (tested via experiments)
""")

    add_heading(doc, '5.5 Campaign Anchor Feature', 2)
    add_para(doc,
        'The most critical feature discovered during this project. The campaign anchor '
        'provides the model with the last known valid analyzer reading as a baseline. '
        'It must be shifted by 1 hour to prevent target leakage:', size=11)
    add_code_block(doc, """
# Step 1: Mark stuck and zero readings as invalid (NaN)
df["C4H8_last_valid"] = df["C4H8_Bottom"].where(
    ~df["C4H8_stuck"] & (df["C4H8_Bottom"] > 0.001),
    other=np.nan
)

# Step 2: Shift by 1 hour (leak-free - anchor at t cannot see y at t)
# Step 3: Forward-fill within each block (up to limit hours)
df["C4H8_campaign_anchor"] = (
    df.groupby("Data_Block")["C4H8_last_valid"]
      .shift(1)
      .ffill(limit=72)  # 72h = 3-day maximum lookback for Model A
)
""")
    add_note_box(doc,
        'WITHOUT the 1-hour shift: anchor[t] = target[t] = direct leakage. '
        'Fake R2 = 0.9997. WITH the shift: anchor[t] = last reading before t. '
        'Real R2 = 0.9074. This difference is the entire foundation of the project.',
        'WARNING')

    add_heading(doc, '5.6 Feature Set Summary', 2)
    rows = [
        ('Tier 1 (Production)', '67', 'Process-only features. No target lags. Safe for deployment.'),
        ('Tier 2 (Research)', '82', 'Includes C4H8/C4H6 lag features. Higher accuracy but analyzer-dependent.'),
        ('Robust 8-Feature', '8', 'Final production set: anchor + 2 ratios + 5 dev24h. Best generalization.'),
    ]
    add_table(doc, ['Feature Set', 'Count', 'Purpose'], rows)

    add_heading(doc, '5.7 Extreme Events Analysis', 2)
    add_para(doc,
        '1,392 rows (12.3%) were found to have been modified by IQR winsorization (not the initially '
        'assumed 114). These rows represent genuine process upsets:', size=11)
    rows = [
        ('Mean C4H8 (extreme rows)', '0.622 wt%', 'vs 0.479 wt% for normal rows'),
        ('Above-spec rate (extreme)', '47.7%', 'vs 38.8% for normal rows'),
        ('Decision', 'KEPT in training', 'Genuine upset behavior, model must learn it'),
    ]
    add_table(doc, ['Metric', 'Value', 'Interpretation'], rows)
    doc.add_page_break()


def section_06_initial_training(doc):
    add_heading(doc, '6. Initial Model Training & The Discovery of Failure (Phase 3)', 1)

    add_heading(doc, '6.1 Training Strategy', 2)
    add_para(doc,
        'Training followed a strict chronological split to simulate real deployment: '
        'Blocks 1-3 as training data, Block 4 as the held-out test set. '
        'The test set (6,081 rows) is larger than the train set (4,332 rows) — this is '
        'intentional as Block 4 represents the full future production campaign.', size=11)

    add_heading(doc, '6.2 Baseline Models (No Feature Engineering)', 2)
    rows = [
        ('Overall Mean Baseline', '-0.003', '0.2199', 'Predict global mean. Useless.'),
        ('Block Mean Baseline', '-0.003', '0.2199', 'Predict per-block mean. Equally useless.'),
        ('Naive Lag-1 (analyzer)', '0.9328', '0.0361', 'Predict C4H8[t] = C4H8[t-1]. Requires working analyzer.'),
    ]
    add_table(doc, ['Baseline', 'R2', 'MAE (wt%)', 'Notes'], rows)

    add_heading(doc, '6.3 Default Models on 67-Feature Set (No Tuning)', 2)
    add_para(doc, 'All four ML models failed catastrophically on the Block 4 test set:', size=11)
    rows = [
        ('LinearRegression', '-3.728', '0.5035', '0.5996', 'No'),
        ('Ridge', '-3.932', '0.5082', '0.6124', 'No'),
        ('RandomForest', '-0.992', '0.2987', '0.3892', 'No'),
        ('XGBoost', '-1.029', '0.3017', '0.3928', 'No'),
    ]
    add_table(doc, ['Model', 'R2', 'MAE', 'RMSE', 'Analyzer?'], rows)
    add_note_box(doc,
        'All R2 values are NEGATIVE. A model that predicts the mean gets R2=0. '
        'Negative R2 means the model is WORSE than just predicting the average. '
        'The model was actively wrong, not just imprecise.',
        'WARNING')

    add_heading(doc, '6.4 The Smoking Gun: Feature Importance Analysis', 2)
    add_para(doc,
        'SHAP and tree feature importance revealed the root cause of failure:', size=11)
    rows = [
        ('1st', 'month_cos', '0.3781', 'CALENDAR FEATURE - year/season proxy!'),
        ('2nd', 'month_sin', '0.0752', 'CALENDAR FEATURE'),
        ('3rd', 'Data_Block', '0.0567', 'CAMPAIGN IDENTIFIER'),
        ('4th', 'Column_Top_Pressure_lag1', '0.0434', 'First real process variable'),
    ]
    add_table(doc, ['Rank', 'Feature', 'Importance', 'Interpretation'], rows)
    add_note_box(doc,
        'The model learned that "what year/month it is" predicts C4H8 better than any '
        'process variable. This is campaign memorization. It performs perfectly on training '
        'data because it memorized historical campaigns, but fails entirely when deployed '
        'in a new campaign (Block 4 of 2025-26).',
        'IMPORTANT')

    add_heading(doc, '6.5 Model B (C4H6) Initial Failure', 2)
    rows = [
        ('LinearRegression', '-305.27', '0.1629', 'Complete failure'),
        ('Ridge', '-298.59', '0.1603', 'Complete failure'),
        ('RandomForest', '-6.25', '0.0190', 'Least bad, still terrible'),
        ('XGBoost', '-34.63', '0.0550', 'Severely negative R2'),
    ]
    add_table(doc, ['Model', 'R2', 'MAE', 'Status'], rows)
    add_para(doc,
        'C4H6 models failed even more catastrophically due to the 24x target collapse '
        '(train mean 0.14 wt% vs test mean 0.0057 wt%). Any model trained on Block 1 data '
        '(mean C4H6 = 0.208 wt%) will predict huge values on Block 4 (mean = 0.0057 wt%).', size=11)
    doc.add_page_break()


def section_07_root_cause(doc):
    add_heading(doc, '7. Root Cause Analysis: Concept Drift & Covariate Shift', 1)

    add_heading(doc, '7.1 Covariate Shift (Input Extrapolation)', 2)
    add_para(doc,
        'Between training (Blocks 1-3) and test (Block 4), the operating envelope shifted '
        'dramatically. Tree models cannot extrapolate beyond training boundaries — they '
        'predict "flat leaf" values for inputs outside training ranges:', size=11)
    rows = [
        ('Reflux_Flow', '[87.89, 105.65]', '[70.42, 105.29]', '38-54% OOB', 'SEVERE'),
        ('Column_Top_Temp', '[29.35, 113.00]', '[18.04, 113.00]', '34% OOB', 'HIGH'),
        ('Reboiling_Steam_Flow', '[18.73, 25.26]', '[14.43, 23.54]', '14% OOB', 'MODERATE'),
    ]
    add_table(doc, ['Feature', 'Train Range', 'Test Range', 'Out-of-Bounds %', 'Severity'], rows)

    add_heading(doc, '7.2 Concept Drift (Correlation Sign Reversal)', 2)
    add_para(doc,
        'The most critical discovery: the DIRECTION of influence of key variables on C4H8 '
        'COMPLETELY REVERSED between training and test periods. This is caused by pressure-'
        'temperature co-dependency in the boiling point equilibrium:', size=11)
    rows = [
        ('Control_Tray_Temp', '-0.367 (negative)', '+0.391 (positive)', 'REVERSED'),
        ('Reflux_Flow', '+0.084 (positive)', '-0.239 (negative)', 'REVERSED'),
        ('Column_Bottom_Temp', '-0.065 (negative)', '+0.223 (positive)', 'REVERSED'),
        ('Temp_Gradient', '-0.254 (negative)', '+0.166 (positive)', 'REVERSED'),
        ('Column_Top_Pressure', '+0.475 (positive)', '+0.058 (near zero)', 'COLLAPSED'),
    ]
    add_table(doc, ['Feature', 'Correlation in Train', 'Correlation in Test (Block 4)', 'Status'], rows)
    add_note_box(doc,
        'The model learned: "Higher control tray temp => lower C4H8" (in 2023-2024). '
        'Block 4 reality: "Higher control tray temp => HIGHER C4H8". '
        'The model was making predictions in the exact wrong direction. '
        'Pearson correlation between predictions and actual = -0.326 (negative!)',
        'IMPORTANT')

    add_heading(doc, '7.3 Physical Explanation: Cascade Control Interference', 2)
    add_para(doc,
        'In normal operation, pressure builds trigger the cascade controller to LOWER '
        'Control_Tray_Temp (automatically). This creates an artificial correlation: '
        'high pressure -> low tray temp -> high C4H8 separation. But in Block 4, '
        'operators changed setpoints, breaking this automated correlation.', size=11)
    add_para(doc, 'This is called "Control Loop Rot" - the model learned an operator '
             'automation artifact as if it were physical separation truth.', bold=True, size=11)
    doc.add_page_break()


def section_08_breakthrough(doc):
    add_heading(doc, '8. The Breakthrough: Campaign Anchor & Feature Ablation', 1)

    add_heading(doc, '8.1 The Experiment Matrix (11 Drift Experiments)', 2)
    add_para(doc,
        'After diagnosing the failure, a systematic series of 11 experiments was run '
        'to find what combination of features could overcome concept drift:', size=11)
    rows = [
        ('Baseline (67 features)', '-0.178', '-1.036', 'month_cos dominates'),
        ('Exp 1: No Calendar', '-0.314', '-0.967', 'Still negative R2'),
        ('Exp 2: No Regime', '-0.208', '-0.882', 'Still negative R2'),
        ('Exp 3: All Removed', '-0.334', '-0.969', 'Still negative R2'),
        ('Exp 4: Pnorm k=3', '-0.268', '-0.730', 'Best among physics attempts'),
        ('Exp 10: Rolling Devs', '-0.326', '-0.852', 'Improved slightly'),
        ('Exp 11: Campaign Anchor', '+0.860', '+0.686', 'FIRST POSITIVE R2. Breakthrough!'),
    ]
    add_table(doc, ['Experiment', 'Pearson', 'R2', 'Notes'], rows)

    add_heading(doc, '8.2 Feature Ablation Study (7 Subsets)', 2)
    add_para(doc,
        'After discovering the anchor works, a systematic ablation study tested '
        'different feature combinations:', size=11)
    rows = [
        ('1. Baseline TIER1', '67', '-0.178', '-1.036', '0.301'),
        ('2. Physics Only (No Temps)', '41', '-0.064', '-0.308', '0.246'),
        ('3. Physics + Pnorm Temps', '59', '-0.220', '-0.585', '0.260'),
        ('4. Deviations & Ratios Only', '7', '-0.077', '-0.319', '0.242'),
        ('5. Physics + Stable + Anchor', '60', '+0.856', '+0.675', '0.107'),
        ('6. Physics Only + Anchor', '42', '+0.832', '+0.659', '0.115'),
        ('7. Deviations & Ratios + Anchor', '8', '+0.922', '+0.835', '0.073'),
    ]
    add_table(doc, ['Subset', 'Features', 'Pearson', 'R2', 'MAE (wt%)'], rows)
    add_note_box(doc,
        'SUBSET 7 is the winner: Only 8 features, but achieves Pearson = 0.922 and R2 = 0.835. '
        'Adding more features (subsets 5, 6) actually REDUCES performance. '
        'Simpler is better — fewer features means less overfitting to campaign-specific noise.',
        'TIP')

    add_heading(doc, '8.3 The 8 Production Features Explained', 2)
    rows = [
        ('C4H8_campaign_anchor', 'Dynamic Calibration', 'Last valid analyzer reading (1h-shifted, 72h limit). The baseline composition anchor.'),
        ('Steam_Feed_Ratio', 'Energy Balance', 'Heat input per unit feed. Dimensionless, campaign-invariant.'),
        ('Reflux_Ratio', 'Mass Balance', 'Cooling input per unit feed. Dimensionless, campaign-invariant.'),
        ('Reboiling_Steam_Flow_dev24h', '24h Deviation', 'Steam change from 24h mean. Captures recent control actions.'),
        ('Reflux_Flow_dev24h', '24h Deviation', 'Reflux change from 24h mean. Captures recent cooling changes.'),
        ('Column_Bottom_Temp_dev24h', '24h Deviation', 'Bottom temp change from 24h mean. Bypasses absolute setpoint drift.'),
        ('Control_Tray_Temp_dev24h', '24h Deviation', 'Tray temp change from 24h mean. Bypasses pressure-temp coupling.'),
        ('Column_Top_Pressure_dev24h', '24h Deviation', 'Pressure deviation from 24h mean. Indicates column loading changes.'),
    ]
    add_table(doc, ['Feature', 'Type', 'Physical Significance'], rows)
    add_para(doc,
        'The key insight: DEVIATIONS from rolling means preserve the physical direction '
        'of correlations across campaigns (steam deviation always correlates positively '
        'with stripping efficiency), while absolute values reverse sign due to setpoint shifts.', size=11)
    doc.add_page_break()


def section_09_model_a(doc):
    add_heading(doc, '9. Model A (C4H8) — Final Development', 1)

    add_heading(doc, '9.1 Model Selection: Why XGBoost?', 2)
    add_para(doc,
        'The robust 8-feature set was tested across all major tree ensemble algorithms:', size=11)
    rows = [
        ('XGBoost (default)', '0.8846', '0.0572'),
        ('CatBoost (Optuna tuned)', '0.9030', '0.0524'),
        ('LightGBM (Optuna tuned)', '0.9147', '0.0494'),
        ('Ensemble (CB + XG + LG)', '0.9052', '0.0513'),
    ]
    add_table(doc, ['Model', 'R2', 'MAE (wt%)'], rows)
    add_para(doc,
        'LightGBM achieved the highest R2 at 0.9147. However, XGBoost was selected for '
        'production deployment after Optuna tuning because: (1) broader industry support, '
        '(2) robust JSON serialization for DCS deployment, (3) Optuna-tuned XGBoost '
        'achieved R2 = 0.9074 (only 0.007 below LightGBM).', size=11, italic=True)

    add_heading(doc, '9.2 Optuna Hyperparameter Optimization', 2)
    add_para(doc,
        '50 Optuna trials were run using 5-fold TimeSeriesSplit cross-validation '
        'on Blocks 1-3 training data. The optimal hyperparameters:', size=11)
    add_code_block(doc, """
Best CV R2 score: 0.7037
Best hyperparameters:
{
    "n_estimators":      102,
    "max_depth":         3,        # KEY: Very shallow trees
    "learning_rate":     0.0405,
    "subsample":         0.8056,
    "colsample_bytree":  0.9360,
    "min_child_weight":  8,
    "gamma":             3.4e-5,
    "reg_alpha":         7.8e-4,
    "reg_lambda":        3.8e-8
}
""")
    add_note_box(doc,
        'max_depth = 3 is extremely shallow. This forces trees to learn simple, monotonic '
        'physical relationships (if steam increases, C4H8 decreases) rather than complex '
        'overfitted splits memorizing campaign-specific correlations.',
        'IMPORTANT')

    add_heading(doc, '9.3 Anchor Lookback Limit: 72h vs 12h Trade-off', 2)
    rows = [
        ('6h', '+0.8450', '0.0694', '90.6%', '93.8%'),
        ('12h', '+0.8450', '0.0694', '91.1%', '94.1%'),
        ('24h', '+0.8341', '0.0726', '92.0%', '94.6%'),
        ('48h', '+0.8341', '0.0726', '93.0%', '95.4%'),
        ('72h', '+0.8345', '0.0726', '94.0%', '96.1%'),
    ]
    add_table(doc, ['Limit', 'R2', 'MAE', 'Train Coverage', 'Test Coverage'], rows)
    add_para(doc,
        'Decision: 72h limit chosen for production. The 12h limit achieves marginally '
        'better R2 (0.8450 vs 0.8345) but drops test coverage to 94.1% (forces 6% fallback time). '
        'The 72h limit increases coverage to 96.1% with negligible performance trade-off.', size=11, italic=True)

    add_heading(doc, '9.4 Final Performance Results', 2)
    rows = [
        ('Block 3 Validation', 'Blocks 1 & 2', 'Block 3', '+0.8848', '0.7694', '0.0817'),
        ('Block 4 Test (FINAL)', 'Blocks 1-3', 'Block 4', '+0.9297', '0.9074', '0.0516'),
    ]
    add_table(doc, ['Split', 'Train', 'Test', 'Pearson', 'R2', 'MAE (wt%)'], rows)

    add_heading(doc, '9.5 Feature Importance (SHAP Gain)', 2)
    rows = [
        ('1', 'C4H8_campaign_anchor', '8.626', '77.8%', 'Composition baseline'),
        ('2', 'Steam_Feed_Ratio', '0.835', '7.5%', 'Heat input ratio'),
        ('3', 'Reboiling_Steam_Flow_dev24h', '0.704', '6.3%', 'Steam change'),
        ('4', 'Control_Tray_Temp_dev24h', '0.611', '5.5%', 'Tray temp change'),
        ('5', 'Column_Top_Pressure_dev24h', '0.426', '3.8%', 'Pressure change'),
        ('6', 'Reflux_Flow_dev24h', '0.384', '3.5%', 'Reflux change'),
        ('7', 'Reflux_Ratio', '0.271', '2.4%', 'Cooling ratio'),
        ('8', 'Column_Bottom_Temp_dev24h', '0.248', '2.2%', 'Bottom temp change'),
    ]
    add_table(doc, ['Rank', 'Feature', 'Gain', 'Relative %', 'Role'], rows)
    add_note_box(doc,
        'The campaign anchor dominates importance (8.626 gain vs 0.835 for 2nd feature). '
        'This is physically correct: current composition depends primarily on recent '
        'measured composition (slow-moving state) plus process corrections.', 'TIP')

    # Model A diagnostic images
    add_heading(doc, '9.6 Model A Diagnostic Plots', 2)
    img1 = os.path.join(IMG_DIR, 'robust_opt_plot_1_actual_vs_predicted.png')
    img2 = os.path.join(IMG_DIR, 'robust_opt_plot_4_residual_vs_time.png')
    img3 = os.path.join(IMG_DIR, 'robust_plot_5_shap.png')

    add_image(doc, img1, 'Figure 1: Model A — Actual vs. Predicted C4H8 on Block 4 Test Set (R2=0.9074)', 5.5)
    add_image(doc, img2, 'Figure 2: Model A — Residual Error vs. Time on Block 4 (Stable, No Drift)', 5.5)
    add_image(doc, img3, 'Figure 3: Model A — SHAP Feature Importance Summary (Anchor Dominates)', 5.5)
    doc.add_page_break()


def section_10_model_b(doc):
    add_heading(doc, '10. Model B (C4H6) — Development & Conclusion', 1)

    add_heading(doc, '10.1 The Target Collapse Problem', 2)
    rows = [
        ('Block 1', '0.208 wt%', 'Cold reboiler. C4H6 extremely high.'),
        ('Block 2', '0.031 wt%', 'Hot regime. C4H6 drops 6.7x from Block 1.'),
        ('Block 3', '0.023 wt%', 'Mixed regime. C4H6 continues declining.'),
        ('Block 4', '0.0057 wt%', 'Efficient fractionation. 37x lower than Block 1!'),
    ]
    add_table(doc, ['Block', 'Mean C4H6', 'Regime'], rows)
    add_note_box(doc,
        'Any ML model trained on Blocks 1-3 (mean C4H6 = 0.1395 wt%) will predict ~0.14 wt% '
        'on Block 4, but actual is 0.0057 wt%. This gives R2 = -17 to -35. '
        'The model is not learning wrong relationships - it is learning right relationships '
        'on the wrong distribution.', 'WARNING')

    add_heading(doc, '10.2 Anchor-Only Baseline Audit', 2)
    add_para(doc,
        'Before attempting ML, an anchor-only audit was run on Block 4:', size=11)
    rows = [
        ('12h Anchor only', '0.9606', '0.0005 (5.5 ppm)', '+0.9830', '98.45%'),
        ('24h Rolling + 12h Anchor', '0.9308', '0.0009 (9.0 ppm)', '+0.9672', '99.56%'),
        ('72h Anchor only', '0.6074', '0.0007 (7.0 ppm)', '+0.7813', '99.76%'),
        ('Block 4 Mean Constant', '0.0000', '0.0032 (32 ppm)', 'Constant', '100%'),
    ]
    add_table(doc, ['Method', 'R2', 'MAE', 'Pearson', 'Coverage'], rows)
    add_note_box(doc,
        'The 12h anchor ALONE achieves R2 = 0.9606 and MAE = 5.5 ppm. '
        'This is BETTER than Model A achieved after all its feature engineering and Optuna tuning. '
        'The C4H6 composition is a slowly-moving state variable that tracks the analyzer directly.',
        'IMPORTANT')

    add_heading(doc, '10.3 The Delta Correction Experiment (ML Makes Things Worse)', 2)
    add_para(doc,
        'A test was run to see if ML could improve upon the anchor by predicting '
        'the high-frequency correction delta:', size=11)
    add_code_block(doc, """
Delta = C4H6_Bottom - C4H6_campaign_anchor_12h

# Model trained to predict Delta, then:
Final_prediction = anchor + predicted_delta
""")
    rows = [
        ('Anchor Only (12h)', '0.9606', '0.000547 (5.5 ppm)', '+0.9830'),
        ('Anchor + XGBoost Delta', '0.9010', '0.001194 (11.9 ppm)', '+0.9499'),
    ]
    add_table(doc, ['Method', 'R2', 'MAE', 'Pearson'], rows)
    add_note_box(doc,
        'Adding ML made performance WORSE. R2 dropped from 0.9606 to 0.9010. '
        'MAE doubled from 5.5 ppm to 11.9 ppm. '
        'The delta model feature importances (all gains < 0.001) confirm that process '
        'variables contain essentially NO additional information beyond the anchor. '
        'Machine learning is mathematically unnecessary for C4H6 prediction.',
        'IMPORTANT')

    add_heading(doc, '10.4 Model B Final Decision: Deterministic Tracker', 2)
    add_bullet(doc, 'Model B is a DETERMINISTIC ANALYZER-TRACKING STATE ESTIMATOR.')
    add_bullet(doc, 'No machine learning. No tree models. No regressors.')
    add_bullet(doc, 'Level 1: Use 12h shifted anchor if available (covers 98.45% of operations).')
    add_bullet(doc, 'Level 2: 24h rolling mean of recent predictions (YELLOW status).')
    add_bullet(doc, 'Level 3: Block 4 historical mean of 0.005663 wt% (RED status, hard timeout > 168h).')

    add_heading(doc, '10.5 Cross-Block Robustness Verification', 2)
    rows = [
        ('Block 2', '0.031 wt%', '0.7518', '42.9 ppm', '+0.8758', '99.71%'),
        ('Block 3', '0.023 wt%', '0.7651', '48.5 ppm', '+0.8821', '99.41%'),
        ('Block 4', '0.0057 wt%', '0.9606', '5.5 ppm', '+0.9830', '98.45%'),
    ]
    add_table(doc, ['Block', 'Mean C4H6', 'R2', 'MAE', 'Pearson', 'Coverage'], rows)
    add_para(doc,
        'The 12h anchor is robust across all blocks. Performance improves dramatically '
        'in Block 4 because the target is nearly stationary (low variance), '
        'making the memory-tracking approach ideal.', size=11, italic=True)
    doc.add_page_break()


def section_11_architecture(doc):
    add_heading(doc, '11. Final Architecture & Dual Soft-Sensor System', 1)

    add_heading(doc, '11.1 System Flow', 2)
    add_code_block(doc, """
EVERY HOUR (t):
================

[1] Retrieve last 24 hours of process data from historian
    Columns: Feed_Flow, Reboiling_Steam_Flow, Reflux_Flow,
             Column_Bottom_Temp, Control_Tray_Temp, Column_Top_Pressure

[2] Retrieve last valid analyzer readings (value + hours_ago)
    - C4H8_Bottom last_valid (within 72h for Model A)
    - C4H6_Bottom last_valid (within 12h for Model B)

[3] MODEL A: XGBoost (C4H8 prediction)
    a. Compute Steam_Feed_Ratio, Reflux_Ratio
    b. Compute dev24h for all 5 process variables
    c. Build 8-feature vector with anchor
    d. Predict via model_A_final_v1.pkl
    e. Return C4H8 prediction + GREEN/YELLOW/RED health

[4] MODEL B: Deterministic Tracker (C4H6 prediction)
    a. If anchor available (< 12h): use anchor directly
    b. If anchor stale (12-168h): use rolling mean prediction
    c. If > 168h offline: use default 0.005663 wt%
    e. Return C4H6 prediction + GREEN/YELLOW/RED health

[5] COMBINE:
    Total_C4 = predicted_C4H8 + predicted_C4H6
    Overall_Health = worst(Health_A, Health_B)
    is_out_of_spec = (Total_C4 > 0.50 wt%)

[6] OUTPUT to DCS/Dashboard
""")

    add_heading(doc, '11.2 Health Status Logic', 2)
    rows = [
        ('GREEN', 'Both analyzers valid and recent', 'Normal operations. Predictions from ML model + fresh anchor.'),
        ('YELLOW', 'At least one analyzer stale (12-72h)', 'Fallback: using 24h rolling mean of recent predictions.'),
        ('RED', 'Analyzer offline > 168h or startup', 'Emergency: using historical campaign mean default values.'),
    ]
    add_table(doc, ['Status', 'Condition', 'What it means'], rows)

    add_heading(doc, '11.3 Safety Ceilings (Non-Negotiable)', 2)
    rows = [
        ('Column_Bottom_Temp', '<= 115.0 deg C', 'Hard alarm/shutdown threshold. Model never recommends exceeding.'),
        ('Column_Top_Pressure', '<= 5.0 kg/cm2g', 'Relief trip threshold. Model never recommends exceeding.'),
        ('Total C4 Alert', '> 0.50 wt%', 'is_out_of_spec flag set to True in output.'),
        ('Optimizer Safety Margin', '<= 0.40 wt%', 'Conservative limit for optimizer (10% below hard spec).'),
    ]
    add_table(doc, ['Parameter', 'Limit', 'Action'], rows)
    doc.add_page_break()


def section_12_inference(doc):
    add_heading(doc, '12. Production Inference & Fallback Logic', 1)

    add_heading(doc, '12.1 Model A Inference Script (predict_c4h8.py)', 2)
    add_code_block(doc, """
def predict_c4h8(process_history, latest_valid_analyzer, previous_predictions):
    # Level 1: Anchor available (within 72 hours)
    if anchor_available (hours_ago <= 72):
        features = {
            "C4H8_campaign_anchor": anchor_value,
            "Steam_Feed_Ratio": steam / feed,
            "Reflux_Ratio": reflux / feed,
            "Reboiling_Steam_Flow_dev24h": steam - mean(steam_24h),
            "Reflux_Flow_dev24h": reflux - mean(reflux_24h),
            "Column_Bottom_Temp_dev24h": bot_temp - mean(bot_temp_24h),
            "Control_Tray_Temp_dev24h": tray_temp - mean(tray_temp_24h),
            "Column_Top_Pressure_dev24h": pressure - mean(pressure_24h)
        }
        pred = model.predict(features)
        return {"predicted_c4h8": pred, "health": "GREEN"}

    # Level 2: Anchor stale, use rolling mean (72h - 168h)
    elif len(previous_preds) >= 6 and not hard_timeout:
        return {"predicted_c4h8": mean(last_24_preds), "health": "YELLOW",
                "fallback_reason": "Analyzer stale >72h"}

    # Level 3: Hard timeout or startup
    else:
        return {"predicted_c4h8": 0.480, "health": "RED",
                "fallback_reason": "Analyzer offline >168h"}
""")

    add_heading(doc, '12.2 Model B Inference Script (predict_c4h6.py)', 2)
    add_code_block(doc, """
DEFAULT_C4H6_MEAN = 0.005663  # Block 4 historical mean

def predict_c4h6(latest_valid_analyzer, previous_predictions):
    # Level 1: Anchor available (12h window)
    if anchor_available (hours_ago <= 12):
        return {"predicted_c4h6": anchor_value, "health": "GREEN"}

    # Level 2: Rolling mean fallback
    elif len(previous_preds) >= 6 and not hard_timeout:
        return {"predicted_c4h6": mean(last_24_preds), "health": "YELLOW",
                "fallback_reason": "Analyzer stale >12h"}

    # Level 3: Default fallback
    else:
        return {"predicted_c4h6": DEFAULT_C4H6_MEAN, "health": "RED",
                "fallback_reason": "Analyzer offline >168h"}
""")

    add_heading(doc, '12.3 Leakage Proof (Programmatic Verification)', 2)
    add_code_block(doc, """
# Formal proof that anchor has no leakage at timestep t:
Original target at t=538: 0.1010
Perturbed target at t=538: 5.1010  # artificially inflated

Original anchor at t=538:     0.1078  # same as before
Perturbed anchor at t=538:    0.1078  # unchanged! No leakage!
Perturbed anchor at t=539:    5.1010  # propagates to NEXT step only

[VERIFIED] Programmatic proof passed: C4H8_campaign_anchor has no current-timestep leakage.
""")
    doc.add_page_break()


def section_13_deployment(doc):
    add_heading(doc, '13. Deployment Guide', 1)

    add_heading(doc, '13.1 Prerequisites', 2)
    add_code_block(doc, """
# Python 3.8+ with dependencies:
pip install pandas numpy xgboost scikit-learn

# Required Python version:
Python 3.10.x (tested)
""")

    add_heading(doc, '13.2 Deployment Directory Structure', 2)
    add_code_block(doc, """
final_v1/
+-- configs/
|   +-- model_A_features.json      (8 production feature names)
|   +-- model_B_features.json      (1 anchor feature name)
+-- models/
|   +-- model_A_final_v1.pkl       (frozen XGBoost binary)
|   +-- model_A_final_v1.json      (frozen XGBoost JSON backup)
+-- inference/
|   +-- predict_c4h8.py            (Model A inference + fallback)
|   +-- predict_c4h6.py            (Model B deterministic tracker)
|   +-- predict_total_c4.py        (Unified combined output)
+-- notebooks/
|   +-- verify_anchor_leakage.py   (Formal leakage proof)
+-- reports/
|   +-- model_A_feature_importance.csv
|   +-- model_A_final_summary.md
+-- README.md
""")

    add_heading(doc, '13.3 DCS Integration Example', 2)
    add_code_block(doc, """
from inference.predict_total_c4 import predict_total_c4
import pandas as pd

# 1. Pull 24h of hourly data from Exaquantum historian
process_data_24h = pd.DataFrame({
    'Feed_Flow':            [80.1, 79.8, ...],  # 24 rows
    'Reboiling_Steam_Flow': [21.0, 20.9, ...],
    'Reflux_Flow':          [88.5, 88.3, ...],
    'Column_Bottom_Temp':   [107.2, 107.1, ...],
    'Control_Tray_Temp':    [72.1, 72.3, ...],
    'Column_Top_Pressure':  [4.05, 4.04, ...]
})

# 2. Latest analyzer readings
latest_c4h8 = {"value": 0.435, "hours_ago": 3}
latest_c4h6 = {"value": 0.0045, "hours_ago": 3}

# 3. Run prediction
result = predict_total_c4(
    process_history=process_data_24h,
    latest_valid_c4h8=latest_c4h8,
    latest_valid_c4h6=latest_c4h6,
    model_a_pkl_path="models/model_A_final_v1.pkl"
)

# 4. Output to DCS
print("Total C4:", result["predicted_total_c4"], "wt%")
print("C4H8:    ", result["predicted_c4h8"], "wt%")
print("C4H6:    ", result["predicted_c4h6"], "wt%")
print("Health:  ", result["prediction_health"])  # GREEN/YELLOW/RED
print("Spec OK: ", not result["is_out_of_spec"])
""")

    add_heading(doc, '13.4 Verification Test', 2)
    add_code_block(doc, """
# Run from project root to verify deployment:
python final_v1/inference/predict_total_c4.py

# Expected output:
=== TESTING predict_total_c4.py ===
Combined Output (Normal Operations):
  Predicted C4H8:     0.4151 wt%
  Predicted C4H6:     0.004500 wt% (45.0 ppm)
  Predicted Total C4: 0.4196 wt%
  Out of Spec (>0.5): False
  Overall Health:     GREEN
  Model A:            Model A (Health: GREEN, Reason: None)
  Model B:            Model B (12h Anchor) (Health: GREEN, Reason: None)
""")
    doc.add_page_break()


def section_14_optimizer(doc):
    add_heading(doc, '14. Future Optimizer Integration', 1)

    add_heading(doc, '14.1 Optimization Objective', 2)
    add_code_block(doc, """
Minimize:  J = Reboiling_Steam_Flow (TPH)   [energy cost reduction]
Subject to:
  Predicted_Total_C4 <= 0.40 wt%  [safety margin below 0.50 spec]
  Column_Bottom_Temp <= 115.0 C    [hard alarm limit]
  Column_Top_Pressure <= 5.0 kg/cm2g  [hard trip limit]
""")

    add_heading(doc, '14.2 Decision Variables & Bounds', 2)
    rows = [
        ('Reboiling_Steam_Flow', '18.0 TPH', '21.0 TPH', '24.4 TPH', '+/- 2.0 TPH/hr'),
        ('Reflux_Flow', '80.0 TPH', '91.1 TPH', '103.9 TPH', '+/- 5.0 TPH/hr'),
    ]
    add_table(doc, ['Variable', 'Min', 'Mean', 'Max', 'Rate Limit'], rows)

    add_heading(doc, '14.3 How the Optimizer Uses the Soft-Sensor', 2)
    add_numbered(doc, 'Generate candidate settings: (+1 TPH reflux, -0.5 TPH steam) etc.')
    add_numbered(doc, 'Compute new Reflux_Ratio and Steam_Feed_Ratio from candidates.')
    add_numbered(doc, 'Re-calculate dev24h deviations with candidate settings.')
    add_numbered(doc, 'Keep anchor values constant (they do not change during optimization step).')
    add_numbered(doc, 'Query predict_total_c4() to get predicted composition.')
    add_numbered(doc, 'Select setting minimizing steam while satisfying constraints.')

    add_heading(doc, '14.4 Self-Correcting via Anchor', 2)
    add_para(doc,
        'When a new analyzer reading arrives and the anchor updates, the optimizer '
        'automatically adjusts recommendations. If feed gets heavier (higher C4H8), '
        'the anchor increases -> optimizer must increase reflux or steam to maintain spec. '
        'This creates a closed-loop self-correcting control system.', size=11)
    doc.add_page_break()


def section_15_files(doc):
    add_heading(doc, '15. File & Folder Structure Reference', 1)

    add_heading(doc, '15.1 Root-Level Files', 2)
    rows = [
        ('data_preprocessing.py', 'Phase 1: Cleans raw Excel, detects stuck readings, assigns blocks, saves parquet.'),
        ('feature_engineering.py', 'Phase 2: Creates 113 engineered features with block-aware lag computations.'),
        ('model_training.py', 'Phase 3: Trains baseline, default, and tier2 models. Produces leaderboard.'),
        ('experiment_catboost.py', 'CatBoost-specific hyperparameter tuning experiments.'),
        ('experiment_optuna.py', 'Initial Optuna tuning experiments on full feature set.'),
        ('generate_documentation.py', 'This script. Generates the comprehensive DOCX documentation.'),
        ('requirements.txt', 'Python package dependencies.'),
        ('readme.txt', 'Project overview and input variable list.'),
    ]
    add_table(doc, ['File', 'Purpose'], rows)

    add_heading(doc, '15.2 data/ Folder', 2)
    rows = [
        ('clean_data.parquet', 'Phase 1 output. 11,343 rows x 24 columns. Cleaned, block-labeled data.'),
        ('features.parquet', 'Phase 2 output. 11,343 rows x 113 columns. All engineered features.'),
    ]
    add_table(doc, ['File', 'Purpose'], rows)

    add_heading(doc, '15.3 notebooks/ Folder (Diagnostic & Analysis Scripts)', 2)
    rows = [
        ('analyze_data.py', 'Initial EDA: correlations, gaps, outliers, stuck readings.'),
        ('analyze_constraints.py', 'Operating limits derivation (P5-P95 ranges, rate constraints).'),
        ('inspect_clean_data.py', 'Visual inspection of Phase 1 parquet output.'),
        ('check_extreme_rows.py', 'Analysis of the 1,392 winsorized extreme event rows.'),
        ('diagnose_predictions.py', 'Out-of-bounds feature detection on test set.'),
        ('experiment_features.py', 'Feature selection ablation study.'),
        ('experiment_regimes.py', 'Hot vs cold regime experiment.'),
        ('inspect_hot_shift.py', 'Feature shift analysis within hot regime.'),
        ('inspect_bias.py', 'Correlation inversion and bias correction analysis.'),
        ('run_drift_experiments.py', 'All 11 drift experiment variants (Pnorm, anchors, deviations).'),
        ('run_feature_ablation_study.py', 'The 7-subset ablation that identified the robust 8-feature set.'),
        ('run_anchor_analysis.py', 'Anchor lookback sensitivity: 6h/12h/24h/48h/72h comparison.'),
        ('run_anchor_tuning.py', 'CatBoost/LightGBM/XGBoost comparison on robust 8 features.'),
        ('tune_robust_xgb.py', '50-trial Optuna tuning for final XGBoost model.'),
        ('verify_anchor_leakage.py', 'Formal programmatic proof of zero leakage in anchor feature.'),
        ('model_b_target_audit.py', 'C4H6 target statistics block-by-block.'),
        ('model_b_anchor_audit.py', 'C4H6 anchor coverage and age analysis on Block 4.'),
        ('anchor_only_baselines.py', 'Pure anchor baseline comparison (12h, 72h, rolling).'),
        ('model_b_delta_model.py', 'The delta correction experiment proving ML degrades C4H6 prediction.'),
        ('model_b_inversion_check.py', 'Cross-block anchor robustness validation.'),
        ('terminal_outputs.md', 'Complete log of all command outputs from every experiment run.'),
    ]
    add_table(doc, ['File', 'Purpose'], rows)

    add_heading(doc, '15.4 experiments/ Folder', 2)
    rows = [
        ('master_leaderboard.csv', 'All experiment results in one table.'),
        ('drift_experiments_summary.csv', 'Results of 11 drift experiments.'),
        ('robust_xgb_optuna_results.json', 'Optuna trial results for XGBoost.'),
        ('diagnostics/', 'All diagnostic PNG plots (actual vs predicted, residuals, SHAP).'),
    ]
    add_table(doc, ['File/Folder', 'Contents'], rows)

    add_heading(doc, '15.5 final_v1/ Folder (Production Release)', 2)
    rows = [
        ('models/model_A_final_v1.json', 'Frozen XGBoost model (JSON format, XGBoost native).'),
        ('models/model_A_final_v1.pkl', 'Frozen XGBoost model (Pickle format, scikit-learn compatible).'),
        ('configs/model_A_features.json', 'Ordered list of 8 production features for Model A.'),
        ('configs/model_B_features.json', 'Feature specification for Model B (anchor-based).'),
        ('inference/predict_c4h8.py', 'Production Model A inference with 3-level fallback.'),
        ('inference/predict_c4h6.py', 'Production Model B deterministic tracker.'),
        ('inference/predict_total_c4.py', 'Unified total C4 prediction with health status.'),
        ('reports/model_A_feature_importance.csv', 'SHAP gain importance for all 8 features.'),
        ('reports/model_A_final_summary.md', 'Comprehensive technical summary report.'),
        ('notebooks/verify_anchor_leakage.py', 'Formal leak-free proof (deployable audit script).'),
    ]
    add_table(doc, ['File', 'Purpose'], rows)
    doc.add_page_break()


def section_16_metrics(doc):
    add_heading(doc, '16. Key Metrics Summary', 1)

    add_heading(doc, '16.1 Model A (C4H8) — Complete Experiment History', 2)
    rows = [
        ('Naive Lag-1 (requires analyzer)', '0.9328', '0.0361', 'Yes', 'Reference baseline'),
        ('Tier 1 XGBoost (67 features)', '-1.029', '0.3017', 'No', 'Initial failure - concept drift'),
        ('Exp 11: Pure Anchor', '0.687', '0.1032', 'No', 'First breakthrough'),
        ('Subset 7: Dev+Ratios+Anchor (8)', '0.835', '0.0726', 'No', 'Ablation winner'),
        ('XGBoost 8-feat (default)', '0.885', '0.0572', 'No', 'Before tuning'),
        ('CatBoost 8-feat (tuned)', '0.903', '0.0524', 'No', 'Alternative'),
        ('LightGBM 8-feat (tuned)', '0.915', '0.0494', 'No', 'Highest R2'),
        ('XGBoost 8-feat (Optuna tuned)', '0.907', '0.0516', 'No', 'FINAL PRODUCTION'),
    ]
    add_table(doc, ['Model', 'R2', 'MAE', 'Analyzer?', 'Notes'], rows)

    add_heading(doc, '16.2 Model B (C4H6) — Key Experiments', 2)
    rows = [
        ('XGBoost Tier 1 (all blocks)', '-34.63', '0.0550', 'No', 'Complete failure'),
        ('XGBoost (Blocks 2+3 only)', '-17 to -35', '~0.05', 'No', 'Still fails'),
        ('12h Anchor alone', '0.9606', '0.0005 (5.5 ppm)', 'Yes (12h)', 'Production choice'),
        ('Anchor + XGBoost Delta', '0.9010', '0.0012 (11.9 ppm)', 'Yes', 'ML made it WORSE'),
    ]
    add_table(doc, ['Model', 'R2', 'MAE', 'Analyzer?', 'Notes'], rows)

    add_heading(doc, '16.3 Final Production Performance', 2)
    rows = [
        ('Model A (C4H8)', 'XGBoost (Optuna, max_depth=3)', '8', '0.9074', '0.0516 wt%', '+0.9297'),
        ('Model B (C4H6)', 'Deterministic Anchor Tracker', '1', '0.9606', '0.0005 wt% (5.5 ppm)', '+0.9830'),
    ]
    add_table(doc, ['Model', 'Algorithm', 'Features', 'R2', 'MAE', 'Pearson'], rows)

    add_heading(doc, '16.4 Key Lessons Learned', 2)
    lessons = [
        ('Leakage kills results', 'anchor[t] = target[t] gave R2=0.9997. anchor[t] = target[t-1] gave R2=0.9074. The 1-hour shift is everything.'),
        ('Fewer features wins', 'Subset 7 (8 features) beats Subset 5 (60 features). Less overfitting on a 4,000-row training set.'),
        ('Deviations beat absolutes', 'dev24h features maintain consistent correlation directions across operating campaigns.'),
        ('ML can be wrong choice', 'Model B proof: pure anchor R2=0.96. Anchor+ML R2=0.90. Sometimes no model IS the model.'),
        ('Physical understanding is essential', 'Concept drift was diagnosed because we understood what cascade controller interference means physically.'),
        ('Test set > train set is fine', '6,081-row test vs 4,332-row train is intentional and should be explained explicitly to reviewers.'),
    ]
    add_table(doc, ['Lesson', 'Explanation'], lessons)
    doc.add_paragraph()

    add_heading(doc, '16.5 Model A Diagnostic Images', 2)
    img1 = os.path.join(FINAL_DIR, 'robust_opt_plot_1_actual_vs_predicted.png')
    img2 = os.path.join(FINAL_DIR, 'robust_opt_plot_4_residual_vs_time.png')
    add_image(doc, img1, 'Final Model A: Actual vs. Predicted on Block 4 (Locked Production Version)', 5.5)
    add_image(doc, img2, 'Final Model A: Residual Stability Over Time (Block 4, No Drift Pattern)', 5.5)

    add_heading(doc, '16.6 Process Regime Scatter', 2)
    regime_img = os.path.join(NB_DIR, 'regime_scatter.png')
    add_image(doc, regime_img, 'Operating Regimes: Reboiler Outlet Temp vs Column Top Temp colored by Total C4 wt%', 5.5)
    doc.add_page_break()

    # Final summary paragraph
    add_heading(doc, 'Closing Summary', 1)
    add_para(doc,
        'This project successfully delivered a production-ready dual soft-sensor system for '
        'the Debutanizer Column. Starting from a raw Excel dataset with significant quality issues, '
        'the team overcame campaign memorization, concept drift, covariate shift, analyzer unreliability, '
        'and a complete target distribution collapse for the C4H6 model.', size=11)
    doc.add_paragraph()
    add_para(doc,
        'The final deliverable is a locked release package (final_v1/) containing frozen model '
        'binaries, inference scripts with 3-level fallback logic, health status indicators for '
        'operator dashboards, and a complete audit trail of all experiments.', size=11)
    doc.add_paragraph()
    rows = [
        ('Model A C4H8 Production R2', '0.9074', 'On Block 4 (8 months of unseen data)'),
        ('Model A C4H8 Production MAE', '0.0516 wt%', 'Well within operational requirements'),
        ('Model B C4H6 Anchor R2', '0.9606', 'Deterministic, no ML needed'),
        ('Model B C4H6 Anchor MAE', '5.5 ppm', 'Sub-ppm class accuracy'),
        ('Features Used (Model A)', '8', 'Down from 67 initial features'),
        ('Analyzer Coverage (Model A)', '96.1%', '72h anchor limit keeps model active'),
        ('Safety Timeout', '168 hours', 'Hard fallback to prevent drift accumulation'),
    ]
    add_table(doc, ['Metric', 'Value', 'Context'], rows)


# ============================================================
# MAIN
# ============================================================

def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Page margins
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    print("Building cover page...")
    build_cover(doc)

    print("Building table of contents...")
    build_toc(doc)

    print("Section 1: Overview...")
    section_01_overview(doc)

    print("Section 2: Process description...")
    section_02_process(doc)

    print("Section 3: Dataset...")
    section_03_dataset(doc)

    print("Section 4: Preprocessing...")
    section_04_preprocessing(doc)

    print("Section 5: Feature engineering...")
    section_05_features(doc)

    print("Section 6: Initial training...")
    section_06_initial_training(doc)

    print("Section 7: Root cause analysis...")
    section_07_root_cause(doc)

    print("Section 8: Breakthrough...")
    section_08_breakthrough(doc)

    print("Section 9: Model A...")
    section_09_model_a(doc)

    print("Section 10: Model B...")
    section_10_model_b(doc)

    print("Section 11: Architecture...")
    section_11_architecture(doc)

    print("Section 12: Inference...")
    section_12_inference(doc)

    print("Section 13: Deployment...")
    section_13_deployment(doc)

    print("Section 14: Optimizer...")
    section_14_optimizer(doc)

    print("Section 15: Files...")
    section_15_files(doc)

    print("Section 16: Metrics...")
    section_16_metrics(doc)

    out_path = os.path.join(BASE_DIR, 'Debutanizer_C4_Project_Documentation.docx')
    doc.save(out_path)
    print(f"\nDocumentation saved to: {out_path}")
    print(f"File size: {os.path.getsize(out_path) / 1024:.1f} KB")


if __name__ == '__main__':
    main()
